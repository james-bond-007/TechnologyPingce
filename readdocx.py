# _*_ coding:utf-8 _*_
# from __future__ import unicode_literals
from docx import Document
import xlwt
import xlrd
from xlutils.copy import copy
import sys
import os
import json
import re


def adddocx():
    file = Document('./表格1.docx')
    tempTable = document.tables
    print("段落数:" + str(len(file.paragraphs)))  # 段落数为13，每个回车隔离一段


# 模板表中每个字段对应的位置，键是字段，值是所在的位置
#dict1 = {}


# 判断是否是英文
def isEnglish(checkStr):
    for ch in checkStr.decode('utf-8'):
        if u'\u4e00' <= ch <= u'\u9fff':
            return False
    return True


# 读取模板表
def readTemplate(filename='./主观评测/常规节目/2020年第一季度常规节目主观评测表（董立坤）20200317.docx'):
    document = Document(filename)  # Document(templete.decode('utf-8'))
    tempTable = document.tables
    table = tempTable[2]
    for table in tempTable:

        rowList = table.rows
        columnList = table.columns
        rowLength = len(rowList)
        columnLength = len(columnList)

        for rowIndex in range(rowLength):
            for columnIndex in range(columnLength):
                cell = table.cell(rowIndex, columnIndex)
                # if isEnglish(cell.text):
                dict1.setdefault(cell.text, [rowIndex, columnIndex])
        print(dict1)


def initdict(documentName):
    document = Document(documentName)
    tempTable = document.tables
    columnIndex = 1
    j = 0
    dict1 = {}
    for table in tempTable:
        for rowIndex in range(len(table.rows) // 8):
            # 生成测试项
            ceshiDict = {}
            for i in range(1, 6):
                ceshiDict.setdefault('测试{0}'.format(i), [])
            ceshiDict.setdefault('综评', [])
            # 节目名称
            cell = table.cell(rowIndex * 8 + 2, columnIndex)
            myname = cell.text
            myname = re.sub('\s', '', myname)
            dict1.setdefault(j, [myname, ceshiDict])
            j += 1
    return dict1


def readdocument(documentname, dict1):
    document = Document(documentname)
    tempTable = document.tables
    columnIndex = 3
    j = 0
    for table in tempTable:
        for rowIndex in range(len(table.rows) // 8):
            tempDicts = dict1[j][1]
            cell = table.cell(rowIndex * 8 + 7, columnIndex)
            myname = cell.text
            myname = re.sub('\s', '', myname)
            if not myname == '':
                tempDicts['综评'].append(myname)
            for i in range(5):
                cell = table.cell(rowIndex * 8 + 2 + i, columnIndex)
                myname = cell.text
                myname = re.sub('\s', '', myname)
                if not myname == '':
                    tempDicts['测试{}'.format(i + 1)].append(myname)
                    if i > 4 or cell.text == table.cell(rowIndex * 8 + 2 + i + 1, columnIndex).text:
                        break
            j += 1
    # print(dict1)
    return dict1


def writeJson(dict1):
    # print(json.dumps(dict1, sort_keys=True, indent=4, ensure_ascii=False))
    jsonData = json.dumps(dict1, sort_keys=True, indent=4, ensure_ascii=False)
    fileObject = open('data.json', 'w')
    fileObject.write(jsonData)
    fileObject.close()


def writeExcel(dict1):
    # 创建excel工作表
    workbook = xlwt.Workbook(encoding='utf-8')
    worksheet = workbook.add_sheet('sheet1')

    # 设置表头
    worksheet.write(0, 0, label='序号')
    worksheet.write(0, 1, label='节目名称')
    worksheet.write(0, 2, label='测试点1')
    worksheet.write(0, 3, label='测试点2')
    worksheet.write(0, 4, label='测试点3')
    worksheet.write(0, 5, label='测试点4')
    worksheet.write(0, 6, label='测试点5')
    worksheet.write(0, 7, label='综评')

    for list_item in range(len(dict1)):
        # print(dict1[list_item])
        worksheet.write(list_item + 1, 0, list_item + 1)
        worksheet.write(list_item + 1, 1, dict1[list_item][0])
        for i in range(1, 6):
            mytext = '。'.join(dict1[list_item][1]['测试{}'.format(i)])
            mytext = mytext.replace("。。", "。")
            mytext = mytext.replace("，。", "。")
            worksheet.write(list_item + 1, i + 1, mytext)
        mytext = '。'.join(dict1[list_item][1]['综评'])
        mytext = mytext.replace("。。", "。")
        mytext = mytext.replace("，。", "。")
        worksheet.write(list_item + 1, 7, mytext)
    # 保存
    workbook.save('./OK.xls')


def writeWord(filename, dict1):
    document = Document(filename)
    tempTable = document.tables
    columnIndex = 3
    j = 0
    for table in tempTable:
        for rowIndex in range(len(table.rows) // 8):
            for i in range(1, 6):
                mytext = '。'.join(dict1[j][1]['测试{}'.format(i)])
                mytext = mytext.replace("。。", "。")
                mytext = mytext.replace("，。", "。")
                table.cell(rowIndex * 8 + 1 + i, columnIndex).text = mytext
            j += 1
    document.save('./主观评测表汇总.docx')
    for table in tempTable:
        for rowIndex in range(len(table.rows) // 8):
            mytext = ''
            for i in range(1, 6):
                mytext += table.cell(rowIndex * 8 + 1 + i, columnIndex).text + '。'
            mytext = mytext.replace("。。", "。")
            mytext = mytext.replace("，。", "。")
            cell1 = table.cell(rowIndex * 8 + 2, columnIndex)
            cell2 = table.cell(rowIndex * 8 + 6, columnIndex)
            mycell = cell1.merge(cell2)
            mycell.text = mytext
    document.save('./主观评测表汇总1.docx')


if __name__ == "__main__":

    documentPath = './主观评测/常规节目'
    documnetsName = os.listdir(documentPath)
    # print(documnetsName)
    documnetsNames = []
    for docxname in documnetsName:
        if docxname[-5:] == '.docx':
            documnetsNames.append(docxname)

    data = initdict(documentPath + '/' + documnetsNames[0])

    for docxname in documnetsNames:
        filename = documentPath + '/' + docxname
        data = readdocument(filename, data)
        print(docxname + '完成了，共计{}个文件，还剩下{}个文件'.format(len(documnetsNames),
                                                       len(documnetsNames) - documnetsNames.index(docxname) - 1))

    # for key, value in dict1.items():
    #    print('{key}:{value}'.format(key=key, value=value))
    writeWord(filename, data)
    writeExcel(data)
    writeJson(data)

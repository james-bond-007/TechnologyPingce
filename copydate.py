# -*- coding: utf-8 -*-
import xdrlib, sys
import xlrd
import xlwt
import xlsxwriter
import easygui
import os
import pandas as pd
import pandas.io.formats.excel
import math
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, Mm
from docx.oxml.ns import qn
import re
from openpyxl import load_workbook


def open_excel(file='test.xls'):
    try:
        data = xlrd.open_workbook(file)
        return data
    except Exception as e:
        print(e)


# 根据索引获取Excel表格中的数据   参数:file：Excel文件路径     colnameindex：表头列名所在行的所以  ，by_index：表的索引
def excel_table_byindex(file='test.xls', colnameindex=0, by_index=0):
    data = open_excel(file)
    table = data.sheets()[by_index]
    nrows = table.nrows  # 行数
    ncols = table.ncols  # 列数
    colnames = table.row_values(colnameindex)  # 某一行数据
    list = []
    for rownum in range(1, nrows):

        row = table.row_values(rownum)
        if row:
            app = {}
            for i in range(len(colnames)):
                app[colnames[i]] = row[i]
            list.append(app)
    return list


# 根据名称获取Excel表格中的数据   参数:file：Excel文件路径     colnameindex：表头列名所在行的所以  ，by_name：Sheet1名称
def excel_table_byname(file='test.xls', colnameindex=0, by_name=u'Sheet1'):
    data = open_excel(file)
    table = data.sheet_by_name(by_name)
    nrows = table.nrows  # 行数
    colnames = table.row_values(colnameindex)  # 某一行数据
    list = []
    for rownum in range(1, nrows):
        row = table.row_values(rownum)
        if row:
            app = {}
            for i in range(len(colnames)):
                app[colnames[i]] = row[i]
            list.append(app)
    return list


def wblist(filedir='./节目单', extension='.xls'):
    '''
    for root, dirs, files in os.walk("./节目单", topdown=False):
        #文件
        for name in files:
            print(os.path.join(root, name))
        #目录
        for name in dirs:
            print(os.path.join(root, name))
    '''
    # 处理文件名
    documnetName = []
    for root, dirs, files in os.walk(filedir, topdown=False):
        # 文件
        for name in files:
            # print(os.path.join(root, name))
            if extension in name:
                documnetName.append([name[0], name[3:-4], root, name])
    documnetName.sort(key=lambda x: x[1])
    # 删除无用列
    documnetNames = []
    for i in documnetName:
        documnetNames.append(i[2:4])
    return documnetNames


def tableHeader(book, documnetName):
    # 处理表头
    # 读表头
    data = open_excel(documnetName)
    table = data.sheets()[0]
    colnames = table.row_values(0)  # 某一行数据
    # 写表头
    pindao = ['卫视', '经济', '都市', '影视', '少儿', '公共', '农民']
    for i in range(7):
        sheet = book.add_sheet(pindao[i], cell_overwrite_ok=True)
        sheet.row(0).write(0, '播出时间')
        for j in range(len(colnames)):
            sheet.row(0).write(j + 1, colnames[j])
            if len(colnames[j]) == 0:
                sheet.col(j + 1).width = 256 * 1
            if colnames[j] == '节目名称':
                sheet.col(j + 1).width = 256 * 40


def copydata(book, root, name, myrownum):
    # 打印文件名称
    print('正在处理：{}'.format(name))
    # 读取源数据
    data = open_excel(os.path.join(root, name))
    table = data.sheets()[0]
    nrows = table.nrows  # 行数
    ncols = table.ncols  # 列数
    colnames = table.row_values(0)  # 某一行数据
    ndate = name.split('-')[2]
    ndate = ndate.split('.')
    ndate = '2020/{}/{}'.format(ndate[0], ndate[1])
    sheetName = int(name.split('-')[0]) - 1
    # print(sheetName)
    # 写入目标文件
    sheet = book.get_sheet(sheetName)
    for rownum in range(1, nrows):
        row = table.row_values(rownum)
        if row:
            sheet.row(myrownum[int(sheetName)]).write(0, ndate)
            for i in range(len(colnames)):
                sheet.row(myrownum[int(sheetName)]).write(i + 1, row[i])
            myrownum[int(sheetName)] += 1


def huizong1():
    # 获取文件名
    documnetNames = wblist()
    # 新建workbook
    book = xlwt.Workbook(encoding='utf-8')
    # 填写表头
    documnetName = os.path.join(documnetNames[0][0], documnetNames[0][1])
    tableHeader(book, documnetName)
    # 各数据表中数据记录初始化
    myrownum = [1, 1, 1, 1, 1, 1, 1]
    # 写数据
    for d in documnetNames:
        copydata(book, d[0], d[1], myrownum)
    print('汇总节目数：{}'.format(myrownum))
    '''
    inMessage = easygui.enterbox(msg='请输入要保存的文件名', title='另存为', default='2020年第一季度节目单汇总')
    if inMessage:
        pass
    else:
        inMessage = 'total'
    '''
    inMessage = 'total'
    book.save('{}.xls'.format(inMessage))


def huizong(filename='./total.xlsx'):
    print('分析目录')
    files = wblist()
    pindao = ['卫视', '经济', '都市', '影视', '少儿', '公共', '农民']
    df = {}
    for i in pindao:
        df[i] = pd.DataFrame()
    for file in files:
        name = file[1]
        print(name)
        ndate = name[2:]
        ndate = ndate.split('.')
        ndate = '2020/{}/{}'.format(ndate[0], ndate[1])
        sheetName = name[:2]
        if sheetName == '科教':
            sheetName = '少儿'
        data = pd.read_excel(os.path.join(file[0], name))
        data.insert(0, '播出时间', ndate)
        if df[sheetName].empty:
            df[sheetName] = data
        else:
            df[sheetName] = pd.concat([df[sheetName], data])
        print(file[1])
    total_number = {}
    with pd.ExcelWriter(filename, engine='xlsxwriter') as writer:
        for i in pindao:
            print('正在写入:{}表'.format(i))
            df[i].to_excel(writer, sheet_name=i, index=False)
            total_number[i] = df[i].shape[0]
    print('各频道节目条数：{}'.format(total_number))


def programefilter(df, jiemu):
    # df = pd.read_excel('total.xls')
    # 清除列
    df = df.loc[:, ~ df.columns.str.contains(':')]
    # df = df.dropna(axis=1, how='all')
    # 清除行
    # 宣
    xdf = df[df['节目名称'].str.contains('中插')]
    # 广告
    blAdvert = df['主视源'].str.contains('广告|广告占位')
    df = df[~ blAdvert]
    # 节目
    jdf = pd.DataFrame()
    for i in jiemu['节目名称']:
        data = df[df['节目名称'].str.contains(i)].copy()
        data['time'] = data['开始时间'].str.split(':').str.get(0).astype(int)
        data = data[data['time']>5]
        jdf = pd.concat([jdf, data])
    # 其它
    blOther = df['节目名称'].str.contains(
        '集|宣|广告|预报|预告|ID|即播|预报|剧情|片头|片尾|剧透|招募|新闻联播|头条|京津冀|这一年|旅游|呼号|导视|多看点|欢乐送|标版|引进节目|专题|logo|LOGO|战略|气象|德龙|专临|前情回顾|先睹为快|公益|年货')
    df = df[~ blOther]

    # print(df.info())
    return df, xdf, jdf


def cleardata(file='filter.xlsx'):
    # Create a Pandas Excel writer using XlsxWriter as the engine.
    with pd.ExcelWriter(file, engine='xlsxwriter') as writer:
        # Convert the dataframe to an XlsxWriter Excel object. Note that we turn off
        # the default header and skip one row to allow us to insert a user defined
        # header.
        workbook = writer.book
        pindao = ['卫视', '经济', '都市', '影视', '少儿', '公共', '农民']
        rownumber = {}
        for i in pindao:
            print('正在整理:{}'.format(i))
            df = pd.read_excel('total.xlsx', i, )
            # 读取频道节目列表
            ndf = pd.read_excel('programlist.xlsx', i, )
            #s = ndf['节目名称'].str.cat(sep='|')
            # 筛选整理节目
            df, xdf, jdf = programefilter(df, ndf)
            rownumber[i] = jdf.shape[0]
            jdf.to_excel(writer, sheet_name='{}'.format(i), startrow=0, na_rep='', index=False)
            worksheet = writer.sheets['{}'.format(i)]
            # 设置节目名称列列宽
            worksheet.set_column(5, 5, 35)
            df.to_excel(writer, sheet_name='{}节目'.format(i), startrow=0, na_rep='', index=False)
            worksheet = writer.sheets['{}节目'.format(i)]
            # 设置节目名称列列宽
            worksheet.set_column(5, 5, 35)
            xdf.to_excel(writer, sheet_name='{}宣'.format(i), startrow=0, na_rep='', index=False)

    print('初选节目数：{}'.format(rownumber))


def readtoPD(file='filter.xlsx'):
    # 读取源数据
    pindao = ['卫视', '经济', '都市', '影视', '少儿', '公共', '农民']
    biaotou = ['播出时间', '节目名称', '长度', '主视源', '磁带条码']
    # 表头格式清除
    # pd.io.formats.excel.header_style = None
    df = pd.read_excel(file, sheet_name=pindao)  # 默认读取第一个sheet，sheet_name=None时，读取所有sheet
    for x in pindao:
        df[x] = df[x].loc[:, biaotou]
        df[x].insert(1, '频道', x)
    result = pd.concat(df, ignore_index=True)
    #
    result['播出时间'].astype('str')
    # 处理节目名称列
    s = result['节目名称']
    s = s.str.replace('^.套', '')
    s = s.str.replace('第20', '_第20')
    s = s.str.replace('-_', '_')
    s = s.str.replace('-2020-', '_2020-')
    d = s.str.split('_', expand=True)
    d.columns = ['节目名称', '期数', '日期']
    s = d['节目名称']
    s = s.str.replace(' 20|-20|-19', '_20', regex=True)
    d['节目名称'] = s.str.split('_').str.get(0).str.strip()
    result.drop(['节目名称'], axis=1, inplace=True)
    result.insert(2, '期数', d['期数'])
    result.insert(2, '节目名称', d['节目名称'])
    # result.insert(1, '日期1', d['日期'])

    result['长度'] = result['长度'].str.split('.').str.get(0)
    result.rename(columns={'主视源': '来源', '磁带条码': '责任人'}, inplace=True)
    return result


def writetoEx1(myData, file='result.xlsx', bPrint=False):
    # 写入精选表
    # Create a Pandas Excel writer using XlsxWriter as the engine.
    with pd.ExcelWriter(file, engine='xlsxwriter') as writer:
        # Convert the dataframe to an XlsxWriter Excel object. Note that we turn off
        # the default header and skip one row to allow us to insert a user defined
        # header.
        myData.to_excel(writer, sheet_name="精选", startrow=0, na_rep='')
        # Get the xlsxwriter workbook and worksheet objects.
        workbook = writer.book
        worksheet = writer.sheets['精选']
        # 加序列号
        worksheet.write(0, 0, '序号')

        # worksheet = workbook.add_worksheet('Sheet1')

        # Add a header format.
        header_format = workbook.add_format({
            'font_size': 12,  # 字体大小
            'bold': 1,  # 是否粗体
            'bg_color': '0f6f32',  # 表格背景颜色
            'font_color': '#E2F3F6',  # 字体颜色
            'align': 'center',  # 对齐方式，left,center,rigth,top,vcenter,bottom,vjustify
            'valign': 'vcenter',  # 垂直居中
            'top': 1,  # 上边框，后面参数是线条宽度
            'left': 1,  # 左边框
            'right': 1,  # 右边框
            'bottom': 1,  # 底边框
            'border_color': 'white',
            'text_wrap': 1,  # 自动换行，可在文本中加 '\n'来控制换行的位置
            # 'num_format':'yyyy-mm-dd' #设定格式为日期格式，如：2017-07-01
        })
        index_format = workbook.add_format({
            'font_size': 12,  # 字体大小
            'bold': 0,  # 是否粗体
            'bg_color': '#0f6f32',  # 表格背景颜色
            'font_color': '#E2F3F6',  # 字体颜色
            'align': 'center',  # 对齐方式，left,center,rigth,top,vcenter,bottom,vjustify
            'valign': 'vcenter',  # 垂直居中
            'top': 1,  # 上边框，后面参数是线条宽度
            'left': 1,  # 左边框
            'right': 1,  # 右边框
            'bottom': 1,  # 底边框
            'border_color': 'white',
            'text_wrap': 1,  # 自动换行，可在文本中加 '\n'来控制换行的位置
            # 'num_format':'yyyy-mm-dd' #设定格式为日期格式，如：2017-07-01
        })
        data_format = workbook.add_format({
            'font_size': 11,  # 字体大小
            'bold': 0,  # 是否粗体
            'bg_color': '#319455',  # 表格背景颜色
            'font_color': '#E2F3F6',  # 字体颜色
            'align': 'center',  # 对齐方式，left,center,rigth,top,vcenter,bottom,vjustify
            'valign': 'vcenter',  # 垂直居中
            'top': 1,  # 上边框，后面参数是线条宽度
            'left': 1,  # 左边框
            'right': 1,  # 右边框
            'bottom': 1,  # 底边框
            'border_color': 'white',
            'text_wrap': 1,  # 自动换行，可在文本中加 '\n'来控制换行的位置
            # 'num_format': 'yyyy-mm-dd'  # 设定格式为日期格式，如：2017-07-01
        })
        data2_format = workbook.add_format({
            'font_size': 11,  # 字体大小
            'bold': 0,  # 是否粗体
            'bg_color': '#54B58A',  # 表格背景颜色
            'font_color': '#E2F3F6',  # 字体颜色
            'align': 'center',  # 对齐方式，left,center,rigth,top,vcenter,bottom,vjustify
            'valign': 'vcenter',  # 垂直居中
            'top': 1,  # 上边框，后面参数是线条宽度
            'left': 1,  # 左边框
            'right': 1,  # 右边框
            'bottom': 1,  # 底边框
            'border_color': 'white',
            'text_wrap': 1,  # 自动换行，可在文本中加 '\n'来控制换行的位置
            # 'num_format': 'yyyy-mm-dd'  # 设定格式为日期格式，如：2017-07-01
        })
        defaul_format = workbook.add_format({
            'font_size': 11,  # 字体大小
            'bold': 0,  # 是否粗体
            'bg_color': 'white',  # 表格背景颜色
            'font_color': 'black',  # 字体颜色
            'align': 'center',  # 对齐方式，left,center,rigth,top,vcenter,bottom,vjustify
            'valign': 'vcenter',  # 垂直居中
            'top': 0,  # 上边框，后面参数是线条宽度
            'left': 0,  # 左边框
            'right': 0,  # 右边框
            'bottom': 0,  # 底边框
            'text_wrap': 1,  # 自动换行，可在文本中加 '\n'来控制换行的位置
            # 'num_format': 'yyyy-mm-dd'  # 设定格式为日期格式，如：2017-07-01
        })
        if bPrint:
            header_format.set_font_color('black')
            header_format.set_fg_color('white')
            header_format.set_border_color('black')
            index_format.set_font_color('black')
            index_format.set_fg_color('white')
            index_format.set_border_color('black')
            data_format.set_font_color('black')
            data_format.set_fg_color('white')
            data_format.set_border_color('black')
            data2_format.set_font_color('black')
            data2_format.set_fg_color('white')
            data2_format.set_border_color('black')

        worksheet.write(0, 0, '序号', header_format)
        # Write the column headers with the defined format.
        colWidth = {'序号': 4.54, '播出时间': 10.21, '频道': 4.54, '节目名称': 29.43, '期数': 11.55,
                    '长度': 8.58, '来源': 6.68, '责任人': 6.68}
        worksheet.set_column(0, 0, colWidth['序号'])
        for col_num, value in enumerate(myData.columns.values):
            # colWidth = max(len(value), max(myData[value].astype(str).str.len()))+3
            # print(colWidth)
            worksheet.set_column(col_num + 1, col_num + 1, colWidth[value])
            worksheet.write(0, col_num + 1, value, header_format)
        for row_num, value in enumerate(myData.index.values):
            # worksheet.set_row(row_num+1, None, data_format)
            worksheet.write(row_num + 1, 0, value + 1, index_format)
        # 列宽
        print('保留节目数：{}'.format(myData.shape[0]))
        worksheet.freeze_panes(1, 1)
        # worksheet.set_column(myData.shape[1]+1, 200, None, defaul_format, {'hidden': 0})
        for i in range(myData.shape[0]):
            if myData.loc[i, '来源'] == '硬盘':
                datastyle = data_format
            else:
                datastyle = data2_format
            for j in range(myData.shape[1]):
                value = str(myData.iloc[i, j])
                if value == 'nan' or value == 'None' or value == '录像机':
                    value = ""
                worksheet.write(i + 1, j + 1, value, datastyle)
                pass


def writetoEx(myData, file='result.xlsx', bPrint=False):
    # 写入精选表
    # Create a Pandas Excel writer using XlsxWriter as the engine.
    with pd.ExcelWriter(file, engine='xlsxwriter') as writer:
        # Convert the dataframe to an XlsxWriter Excel object. Note that we turn off
        # the default header and skip one row to allow us to insert a user defined
        # header.
        myData.to_excel(writer, sheet_name="精选", startrow=0, na_rep='')
        # Get the xlsxwriter workbook and worksheet objects.
        workbook = writer.book
        worksheet = writer.sheets['精选']
        # 加序列号
        worksheet.write(0, 0, '序号')

        # worksheet = workbook.add_worksheet('Sheet1')

        # Add a header format.
        header_format = workbook.add_format({
            'font_size': 12,  # 字体大小
            'bold': 1,  # 是否粗体
            'bg_color': '0f6f32',  # 表格背景颜色
            'font_color': '#E2F3F6',  # 字体颜色
            'align': 'center',  # 对齐方式，left,center,rigth,top,vcenter,bottom,vjustify
            'valign': 'vcenter',  # 垂直居中
            'top': 1,  # 上边框，后面参数是线条宽度
            'left': 1,  # 左边框
            'right': 1,  # 右边框
            'bottom': 1,  # 底边框
            'border_color': 'white',
            'text_wrap': 1,  # 自动换行，可在文本中加 '\n'来控制换行的位置
            # 'num_format':'yyyy-mm-dd' #设定格式为日期格式，如：2017-07-01
        })
        index_format = workbook.add_format({
            'font_size': 12,  # 字体大小
            'bold': 0,  # 是否粗体
            'bg_color': '#0f6f32',  # 表格背景颜色
            'font_color': '#E2F3F6',  # 字体颜色
            'align': 'center',  # 对齐方式，left,center,rigth,top,vcenter,bottom,vjustify
            'valign': 'vcenter',  # 垂直居中
            'top': 1,  # 上边框，后面参数是线条宽度
            'left': 1,  # 左边框
            'right': 1,  # 右边框
            'bottom': 1,  # 底边框
            'border_color': 'white',
            'text_wrap': 1,  # 自动换行，可在文本中加 '\n'来控制换行的位置
            # 'num_format':'yyyy-mm-dd' #设定格式为日期格式，如：2017-07-01
        })
        data_format = workbook.add_format({
            'font_size': 11,  # 字体大小
            'bold': 0,  # 是否粗体
            'bg_color': '#319455',  # 表格背景颜色
            'font_color': '#E2F3F6',  # 字体颜色
            'align': 'center',  # 对齐方式，left,center,rigth,top,vcenter,bottom,vjustify
            'valign': 'vcenter',  # 垂直居中
            'top': 1,  # 上边框，后面参数是线条宽度
            'left': 1,  # 左边框
            'right': 1,  # 右边框
            'bottom': 1,  # 底边框
            'border_color': 'white',
            'text_wrap': 1,  # 自动换行，可在文本中加 '\n'来控制换行的位置
            # 'num_format': 'yyyy-mm-dd'  # 设定格式为日期格式，如：2017-07-01
        })
        data2_format = workbook.add_format({
            'font_size': 11,  # 字体大小
            'bold': 0,  # 是否粗体
            'bg_color': '#54B58A',  # 表格背景颜色
            'font_color': '#E2F3F6',  # 字体颜色
            'align': 'center',  # 对齐方式，left,center,rigth,top,vcenter,bottom,vjustify
            'valign': 'vcenter',  # 垂直居中
            'top': 1,  # 上边框，后面参数是线条宽度
            'left': 1,  # 左边框
            'right': 1,  # 右边框
            'bottom': 1,  # 底边框
            'border_color': 'white',
            'text_wrap': 1,  # 自动换行，可在文本中加 '\n'来控制换行的位置
            # 'num_format': 'yyyy-mm-dd'  # 设定格式为日期格式，如：2017-07-01
        })
        defaul_format = workbook.add_format({
            'font_size': 11,  # 字体大小
            'bold': 0,  # 是否粗体
            'bg_color': 'white',  # 表格背景颜色
            'font_color': 'black',  # 字体颜色
            'align': 'center',  # 对齐方式，left,center,rigth,top,vcenter,bottom,vjustify
            'valign': 'vcenter',  # 垂直居中
            'top': 0,  # 上边框，后面参数是线条宽度
            'left': 0,  # 左边框
            'right': 0,  # 右边框
            'bottom': 0,  # 底边框
            'text_wrap': 1,  # 自动换行，可在文本中加 '\n'来控制换行的位置
            # 'num_format': 'yyyy-mm-dd'  # 设定格式为日期格式，如：2017-07-01
        })
        if bPrint:
            header_format.set_font_color('black')
            header_format.set_fg_color('white')
            header_format.set_border_color('black')
            index_format.set_font_color('black')
            index_format.set_fg_color('white')
            index_format.set_border_color('black')
            data_format.set_font_color('black')
            data_format.set_fg_color('white')
            data_format.set_border_color('black')
            data2_format.set_font_color('black')
            data2_format.set_fg_color('white')
            data2_format.set_border_color('black')

        worksheet.write(0, 0, '序号', header_format)
        # Write the column headers with the defined format.
        colWidth = {'序号': 4.54, '播出时间': 10.21, '频道': 4.54, '节目名称': 29.43, '期数': 11.55,
                    '长度': 8.58, '来源': 6.68, '责任人': 6.68}
        worksheet.set_column(0, 0, colWidth['序号'])
        for col_num, value in enumerate(myData.columns.values):
            # colWidth = max(len(value), max(myData[value].astype(str).str.len()))+3
            # print(colWidth)
            worksheet.set_column(col_num + 1, col_num + 1, colWidth[value])
            worksheet.write(0, col_num + 1, value, header_format)
        for row_num, value in enumerate(myData.index.values):
            # worksheet.set_row(row_num+1, None, data_format)
            worksheet.write(row_num + 1, 0, value + 1, index_format)
        # 列宽
        print('保留节目数：{}'.format(myData.shape[0]))
        worksheet.freeze_panes(1, 1)
        # worksheet.set_column(myData.shape[1]+1, 200, None, defaul_format, {'hidden': 0})
        for i in range(myData.shape[0]):
            if myData.loc[i, '来源'] == '硬盘':
                datastyle = data_format
            else:
                datastyle = data2_format
            for j in range(myData.shape[1]):
                value = str(myData.iloc[i, j])
                if value == 'nan' or value == 'None' or value == '录像机':
                    value = ""
                worksheet.write(i + 1, j + 1, value, datastyle)
                pass


def main():
    tables = excel_table_byindex()
    for row in tables[-5:]:
        print(row)

    tables = excel_table_byname()
    for row in tables[-5:]:
        print(row)


def cleardocument(documentname='常规节目主观评测表（评委）.docx'):
    document = Document(documentname)
    tempTable = document.tables
    for table in tempTable:
        for rowIndex in range(2, 20, 8):
            # tempdf = df.loc[j]
            for i in range(3):
                cell = table.cell(rowIndex, i)
                # print(cell.text)
                cell.text = ''
    document.save('常规节目主观评测表（评委）.docx')


def writedocument(file='freeze.xlsx', sheet='常规', blMerge=False):
    df = pd.read_excel(file, sheet_name=sheet)
    df = df.loc[:, ['序号', '节目名称']]
    df = df.dropna(how='all')
    df['序号'] = df['序号'].astype("int")

    document = Document()

    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    sections = document.sections
    section = sections[0]
    section.left_margin = Mm(19.1)
    section.right_margin = Mm(19.1)

    for j in range(df.shape[0]):
        if j % 3 == 0:
            p = document.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run('第十三期电视节目技术质量主观评测打分表')
            r.font.size = Pt(15)
            r.bold = True
        tempTable = bulitTable(document, blMerge)
        for i in range(df.shape[1]):
            cell = tempTable.cell(2, i)
            cell.text = str(df.iloc[j, i])
            # 设置中间对齐
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        print(j)
    document.save(sheet + '模板.docx')
    return sheet + '模板.docx'


def bulitTable(document, blMerge=False):
    table = document.add_table(rows=8, cols=14, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style.font.size = Pt(10.5)  # Pt(10.5)
    table.autofit = False

    for i in range(14):
        table.columns[i].width = Mm(11.5)
    table.columns[1].width = Mm(18.7)
    table.columns[2].width = Mm(15.7)
    table.columns[13].width = Mm(15.7)
    for i in range(8):
        table.rows[i].height = Mm(8.5)
    table.rows[1].height = Mm(11.4)

    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 1).merge(table.cell(1, 1))
    table.cell(0, 2).merge(table.cell(1, 2))
    table.cell(0, 3).merge(table.cell(0, 9))
    table.cell(0, 10).merge(table.cell(0, 12))
    table.cell(0, 13).merge(table.cell(1, 13))
    table.cell(2, 0).merge(table.cell(7, 0))
    table.cell(2, 1).merge(table.cell(7, 1))
    # 总评分设置格式
    cell1 = table.cell(2, 13).merge(table.cell(7, 13))
    cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 测试点及综评
    for i in range(2, 8):
        cell1 = table.cell(i, 3).merge(table.cell(i, 12))
        # cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if blMerge:
        table.cell(2, 2).merge(table.cell(7, 2))
        table.cell(2, 3).merge(table.cell(7, 12))

    title1 = ['序号', '节目名称', '测试点', '图像（70）分', '声音（30）分', '总评分']
    j = 0
    for i in [0, 1, 2, 3, 10, 13]:
        cell1 = table.cell(0, i)
        cell1.text = title1[j]
        cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        j += 1
    title2 = ['杂波\r干扰', '清晰度', '亮度\r层次', '色彩\r保真', '制作\r难度',
              '素材\r资料', '灯光\r舞美', '声音\r质量', '声音\r音量', '声画\r协调']
    for i in range(len(title2)):
        cell1 = table.cell(1, i + 3)
        cell1.text = title2[i]
        cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if blMerge:
        table.cell(2, 2).text = '综评'
        table.cell(2, 2).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(2, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    else:
        table.cell(7, 2).text = '综评'
        table.cell(7, 2).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(7, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for i in range(5):
            table.cell(i + 2, 2).text = '测试{}'.format(i + 1)
            table.cell(i + 2, 2).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(i + 2, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return table


def tongjifen(file='主观分汇总.xlsx', blsort=True, precision=0):
    # pd.set_option('precision', 3)
    # precision = 0  # 设定保留小数位
    documnents = wblist(filedir='./主观评测/常规节目', extension='.docx')
    data = {}
    zhuanjia = []
    for i in documnents:
        name = re.findall(r'[（](.*?)[）]', i[1])
        print(name[0])
        data['序号'] = []
        data['节目名称'] = []
        data[name[0]] = []
        zhuanjia.append(name[0])
        filename = os.path.join(i[0], i[1])
        documnent = Document(filename)
        tables = documnent.tables
        for table in tables:
            for j in range(len(table.rows) // 8):
                data['序号'].append(table.cell(j * 8 + 2, 0).text)
                data['节目名称'].append(table.cell(j * 8 + 2, 1).text)
                fen = table.cell(j * 8 + 2, 13).text
                # 去掉空格
                # str_list = fen.split()
                # fen = ''.join(str_list)

                fen = fen.strip()
                if len(fen) > 2:
                    print('--->', table.cell(j * 8 + 2, 1).text, fen)
                # if len(fen) > 1:
                    fen = re.findall("\d+", fen)[0]
                elif len(fen) < 2:
                    print('--->', table.cell(j * 8 + 2, 1).text, fen)
                data[name[0]].append(fen)
    # print(data)
    df = pd.DataFrame.from_dict(data)
    df[zhuanjia] = df[zhuanjia].apply(pd.to_numeric)
    temp = df[zhuanjia]
    df['主观'] = temp.mean(axis=1)
    # df['avg'] = df['avg'].round(2)
    if blsort:
        df = df.sort_values(by=['主观', '序号'], ascending=[False, True])
        df.reset_index(drop=True, inplace=True)
    print(df.info())

    with pd.ExcelWriter(file, engine='xlsxwriter') as writer:
        df.to_excel(writer, sheet_name="汇总", startrow=0, na_rep='', index=False)
        # 把公式写入指定单元格
        worksheet = writer.sheets['汇总']
        n = 1 #写入公式列是倒数第几列
        col_num = df.shape[1] - n + ord('A')
        for i in range(df.shape[0]):
            worksheet.write_formula('{0}{1}'.format(chr(col_num), i + 2),
                    '=ROUND(AVERAGE({0}{2}:{1}{2}),{3})'.format(chr(col_num-9), chr(col_num-1), i + 2, precision))

def huizongfen(file='总分统计表.xlsx', precision=0):
    # precision = 0  # 设定保留小数位
    df1 = pd.read_excel('./主观分汇总.xlsx')
    # print(df1)
    df2 = pd.read_excel('./客观分.xlsx', header=1)
    df2 = df2[['节目名称', '总评分']]
    df2 = df2.dropna(how='all')
    df2.rename(columns={'总评分': '客观'}, inplace=True)
    # print(df2)
    # result = pd.concat([df1, df2], axis=1, join_axes=[df1.index])
    result = pd.merge(df1, df2, on='节目名称')
    print(result.info())
    result['总分'] = result.apply(lambda x: (x['主观'] * 9 + x['客观']) / 10, axis=1)
    result['总分'] = result['总分'].round(precision)
    result['等级'] = result.apply(
        lambda x: '优秀' if x['总分'] >= 90 else
        '良好' if x['总分'] >= 85 else
        '良' if x['总分'] >= 80 else
        '及格' if x['总分'] >= 60 else '不及格', axis=1
    )
    # result = result.sort_values(by=['总分', '主观', '客观', '序号'],
    #                             ascending=[False, False, False, True])
    # result.reset_index(drop=True, inplace=True)

    with pd.ExcelWriter(file, engine='xlsxwriter') as writer:
        result.to_excel(writer, sheet_name="汇总", startrow=0, na_rep='', index=False)
        # 把公式写入指定单元格
        worksheet = writer.sheets['汇总']
        n = 2  # 写入公式列是倒数第几列
        col_num = result.shape[1] - n + ord('A')
        for i in range(result.shape[0]):
            worksheet.write_formula('{0}{1}'.format(chr(col_num), i + 2),
                    '=ROUND(({0}{2}*9+{1}{2})/10,{3})'.format(chr(col_num-2), chr(col_num-1), i + 2, precision))
            worksheet.write_formula('{0}{1}'.format(chr(col_num+1), i + 2),
                    '=IF({0}{1}<70,"不及格",IF({0}{1}<80,"及格",IF({0}{1}<85,"良",IF({0}{1}<90,"良好","优秀"))))'.format(
                        chr(col_num), i + 2))


def jisuanfen(file='总分统计表.xlsx'):
    df = pd.read_excel(file)
    df = df.sort_values(by=['总分', '主观', '客观', '序号'], ascending=[False, False, False, True])
    df.reset_index(drop=True, inplace=True)
    print(df)
    # 在原文件中增加新表
    # book = load_workbook(file)
    # with pd.ExcelWriter(file, engine='openpyxl') as writer:
    #     writer.book = book
    #     writer.sheets = dict((ws.title, ws) for ws in book.worksheets)
    #     df.to_excel(writer, sheet_name="排序", startrow=1, na_rep='')
    # 写入格式
    with pd.ExcelWriter(file, engine='xlsxwriter') as writer:
        df.to_excel(writer, sheet_name="排序", startrow=1, na_rep='')
        # 写入格式
        workbook = writer.book
        worksheet = writer.sheets['排序']
        merge_format = workbook.add_format({
            'bold': True,
            #'border': 6,
            'font_size': 17,
            'align': 'center',  # 水平居中
            'valign': 'vcenter',  # 垂直居中
            #'fg_color': '#D7E4BC',  # 颜色填充
        })
        header_format = workbook.add_format({
            'bold': True,
            'border': 1,
            'font_size': 11,
            'align': 'center',  # 水平居中
            'valign': 'vcenter',  # 垂直居中
            'fg_color': '#D7E4BC',  # 颜色填充
        })
        data_format = workbook.add_format({
            # 'bold': True,
            'border': 1,
            'font_size': 11,
            'align': 'center',  # 水平居中
            'valign': 'vcenter',  # 垂直居中
            # 'fg_color': '#D7E4BC',  # 颜色填充
            'text_wrap': 1,
        })
        data = '2020年第二季度常规节目评测分数统计排序表'
        worksheet.merge_range(0, 0, 0, df.shape[1], data, merge_format)
        #worksheet.write(0, 0, '2020年第一季度常规节目评测分数统计排序表')
        worksheet.write(1, 0, '排名', header_format)
        # Write the column headers with the defined format.
        colWidth = {'排名':4.93, '序号': 4.93, '节目名称': 27.93, '专家': 6.93,
                    '主客观分': 8.58, '总分': 4.93}
        worksheet.set_column(0, 1, colWidth['排名'])
        worksheet.set_column(2, 2, colWidth['节目名称'])
        worksheet.set_column(3, 11, colWidth['专家'])
        worksheet.set_column(12, 15, colWidth['总分'])
        for col_num, value in enumerate(df.columns.values):
            # colWidth = max(len(value), max(myData[value].astype(str).str.len()))+3
            # print(colWidth)
            #worksheet.set_column(col_num + 1, col_num + 1, colWidth[value])
            worksheet.write(1, col_num + 1, value, header_format)
        #排名
        for row_num, value in enumerate(df.index.values):
            # worksheet.set_row(row_num+1, None, data_format)
            worksheet.write(row_num + 2, 0, value + 1, header_format)
        # 列宽
        # worksheet.set_column(myData.shape[1]+1, 200, None, defaul_format, {'hidden': 0})
        for i in range(df.shape[0]):
            for j in range(df.shape[1]):
                value = df.iloc[i, j]
                worksheet.write(i + 2, j + 1, value, data_format)

        # 把公式写入指定单元格
        worksheet = writer.sheets['排序']
        precision = 0  # 设定保留小数位
        firstrow = 3
        for i in range(df.shape[0]):
            worksheet.write_formula('O{}'.format(i + firstrow),
                                    '=ROUND(M{0}*0.9+N{0}*0.1,{1})'.format(i + firstrow, precision),
                                    data_format)
            # worksheet.write_formula('A{}'.format(i + firstrow),
            #                         '=RANK(O{0},O{0}:O{1}))'.format(i + firstrow, df.shape[0]),
            #                         data_format)


def database(file='database.xlsx'):
    df1 = pd.read_excel('freeze.xlsx', sheet_name='常规', parse_dates=['播出时间'])
    df1['播出时间'] = df1['播出时间'].dt.strftime('%Y/%m/%d')

    df1.insert(3, '首播频道', df1['频道'])
    # 替换频道名称
    pindao = {'卫视': '河北卫视', '经济': '河北经济', '都市': '河北都市', '影视': '河北影视', '少儿': '河北少儿', '公共': '河北公共'}
    df1['首播频道'].replace(pindao, inplace=True)

    df1.insert(8, '节目来源', df1['来源'])
    # 替换来源
    laiyuan = {'硬盘': '备播系统', '自送1': '大洋技审'}
    df1['节目来源'].replace(laiyuan, inplace=True)
    # 替换磁带库
    bool = df1['节目来源'].str.contains('P')
    filter_data = df1['节目来源'][bool]
    # print(filter_data.to_list())
    df1['节目来源'].replace(filter_data.to_list(), '磁带库', inplace=True)
    # print(df1['来源'])

    # print(df1.info())
    df2 = pd.read_excel('总分统计表.xlsx', sheet_name='排序', header=1)
    # df2 = df2[['节目名称', '主观分', '客观分', '总分']]
    # print(df2.info())
    df = pd.merge(df1, df2)
    print(df.head(5), df.info())
    # 排序
    df = df.sort_values(by=['排名'], ascending=True)
    # 录制地点分类
    introduce = ['引进节目', '引进包装']
    outdoor = ['外景录制']
    tai_outer = ['台外录制', '联合录制', '北京录制', '西院五楼新媒体演播室', '北京']
    tai_inter = ['800演播室', '400演播室', '300演播室', '260演播室', '120演播室', '110演播室', '70演播室', ]
    # 写入录制地点
    df['地点'] = df.apply(
        lambda x: '台内录制' if x['录制地点'] in tai_inter else
        '台外录制' if x['录制地点'] in tai_outer else
        '引进节目' if x['录制地点'] in introduce else
        '外景录制' if x['录制地点'] in outdoor else None, axis=1
    )

    # 制作方式分类
    tai_outer = ['北京制作', '台外制作', '联合制作']
    tai_inter = ['包装制作', '高清自制', '直播', '120自制', '影视自制', '少儿自制', '农民自制', '广告自制', '录播']
    # 写入录制地点
    df['方式'] = df.apply(
        lambda x: '台内制作' if x['制作方式'] in tai_inter else
        '台外制作' if x['制作方式'] in tai_outer else None, axis=1
    )
    # 增加评语列
    df['评语'] = ''

    with pd.ExcelWriter(file, engine='xlsxwriter') as writer:
        df.to_excel(writer, sheet_name="基础", na_rep='', index=False)


if __name__ == "__main__":
    # 第一步
    # 将节目单汇总成各频道表
    # huizong()
    # 删除无用列，整理广告等无用节目
    # cleardata()

    # 第二步
    # 把各频道筛选出来的节目汇总到一张表中，处理节目名称
    # writetoEx(readtoPD(), bPrint=False)
    # 第三步
    # 生成主观评测表
    # writedocument(sheet='常规', blMerge=False)
    # 第四步
    # 汇总主观评测分数
    # tongjifen(blsort=False)
    # 计算总分
    # huizongfen()
    # jisuanfen()
    # 合成数据基础表
    database()

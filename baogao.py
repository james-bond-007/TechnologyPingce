import xlsxwriter
from docx import Document
from docx.oxml import parse_xml
from docx.shared import Inches, Pt, Mm
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
import re
import pandas as pd
import numpy as np
from openpyxl import load_workbook
from word_tool import plt_bar, plt_plot, plt_pie, plt_scatter
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm


def initBaogao(file='./demo.docx'):
    document = Document()
    # 设置一个空白样式
    style = document.styles['Normal']
    # 设置西文字体
    style.font.name = u'宋体'  # 'Times New Roman'
    # 设置中文字体
    style.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # '微软雅黑')#

    style = document.styles['Heading 0']
    font = style.font
    # 获取段落样式
    paragraph_format = style.paragraph_format
    # 首行缩进0.74厘米，即2个字符
    paragraph_format.first_line_indent = Mm(7.4)

    # 标题级别
    heading1 = 0
    heading2 = 1
    sections = document.sections
    current_section = sections[-1]
    # 第一章
    p = document.add_heading('第一章 本期评测综述', heading1)
    '''
    p1 = document.add_paragraph()
    run = p1.add_run(u'第一章 本期评测综述')
    run.font.name = u'宋体'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    '''
    p = document.add_paragraph('A plain paragraph having some ')
    p.paragraph_format.first_line_indent = Mm(7.4)
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True
    heading_first = ['一、参评节目数量',
                     '二、综合得分等级及占比',
                     '三、技术质量优秀节目、不达标节目列表',
                     '1、优秀节目',
                     '2、不达标节目',
                     '四、本期亮点',
                     '五、本期报告用语说明', ]
    for heading in heading_first:
        head = document.add_heading(heading, heading2)
        # head.element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        p = document.add_paragraph('please input some words.')
        p.paragraph_format.first_line_indent = Mm(7.4)
    p_shuoming = ['台外录制：以台外演播室录制为主，包括北京的演播室、联合录制、西院五楼新媒体演播室',
                  '台内录制：以台内演播室录制为主，包括800演播厅、4heading10演播厅、300演播厅、260演播厅、120演播室、110演播室、70演播室',
                  '台外制作：是指在台外制作、但不包括在北京制作',
                  '包装制作：在我台云平台高清制作网、由电视制作中心包装制作人员完成',
                  '高清自制：在我台云平台高清制作网、由编辑人员自行完成',
                  '120自制：在电视制作中心120演播室、由编辑人员自行完成',
                  '影视自制：在影视频道制作网、由编辑人员自行完成',
                  '少儿自制：在少儿频道制作网、由编辑人员自行完成',
                  '农民自制：在农民制作网、由编辑人员自行完成',
                  '广告自制：在广告发展公司制作']
    for p in p_shuoming:
        i = document.add_paragraph(p)
        i.paragraph_format.first_line_indent = Mm(7.4)
    # 第二章
    # document.add_page_break()
    document.add_section()
    document.add_heading('第二章 节目技术质量评测结果', heading1)
    heading_first = ['一、参评节目',
                     '二、综合得分排序', ]
    for heading in heading_first:
        document.add_heading(heading, heading2)
        p = document.add_paragraph('please input some words.')
        p.paragraph_format.first_line_indent = Mm(7.4)
    # 第三章
    document.add_section()
    document.add_heading('第三章  数据分析', heading1)
    heading_first = ['一、按频道分析',
                     '频道分析表',
                     '频道分析图',
                     '二、按录制地点分析',
                     '录制地点分析表',
                     '录制地点分析图',
                     '三、按制作方式分析',
                     '制作方式分析表',
                     '制作方式分析图',
                     '专家意见及建议']
    for heading in heading_first:
        document.add_heading(heading, heading2)
        p = document.add_paragraph('please input some words.')
        p.paragraph_format.first_line_indent = Mm(7.4)
    document.save(file)


def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def setCellBackgroundColor(self, cell, rgbColor):
    if not isinstance(rgbColor, RGBValue):
        print('rgbColor is not RGBValue...', type(rgbColor))
        return
    hr = str(hex(int(rgbColor.r)))[-2:]
    hg = str(hex(int(rgbColor.g)))[-2:]
    hb = str(hex(int(rgbColor.b)))[-2:]
    colorStr = hr + hg + hb
    # print(colorStr)
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value=colorStr))
    cell._tc.get_or_add_tcPr().append(shading_elm_1)


def canping_program(file='demo.docx'):
    df = pd.read_excel('database.xlsx')
    df = df[['序号', '节目名称', '播出时间']]
    inRow = df.shape[0] // 2 + 1
    document = Document(file)
    # 将表格插入指定位置
    for p in document.paragraphs:
        if re.match("^Heading \d+$", p.style.name):
            if p.text == '一、参评节目':
                print(p.text)
                table = document.add_table(rows=inRow, cols=7, style='Table Grid')
                move_table_after(table, p)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, 3).merge(table.cell(inRow - 1, 3))
    table.autofit = False
    table.columns[0].width = Mm(12)
    table.columns[1].width = Mm(42)
    table.columns[2].width = Mm(24)
    table.columns[3].width = Mm(4)
    table.columns[4].width = Mm(12)
    table.columns[5].width = Mm(42)
    table.columns[6].width = Mm(24)
    columns = df.columns.to_list()
    for i in range(2):
        for column in columns:
            cell1 = table.cell(0, i * 4 + columns.index(column))
            cell1.text = column
            cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # 设置表头底色
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value='#8DB4E2'))
            cell1._tc.get_or_add_tcPr().append(shading_elm_1)
    # 让正确转行
    inRow -= 1
    table.rows[0].height = Mm(7.2)  # 表头行高
    for index, row in df.iterrows():
        table.rows[index % inRow + 1].height = Mm(7.2)  # 数据行高
        for i in range(3):
            cell1 = table.cell(index % inRow + 1, index // inRow * 4 + i)
            cell1.text = str(row[columns[i]])
            cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    document.save(file)


def fenxi_zhanbi(file='demo.docx'):
    df = pd.read_excel('database.xlsx')
    df = df[['频道', '节目名称', '播出时间', '录制地点',
             '制作方式', '制片人', '主观', '客观', '总分', '等级']]
    df['主观'] = df['主观'].round(0)
    df['客观'] = df['客观'].astype(np.int64)
    df['总分'] = df['总分'].astype(np.int64)

    pindao = ['卫视', '经济', '都市', '影视', '少儿', '公共', '农民']
    fenxi = []
    for i in pindao:
        # 按频道筛选
        df_temp = df[df['频道'] == i]
        # 统计等级个数
        temp = df_temp.loc[:, '等级'].value_counts()
        # 用频道名称重新命名序列名
        temp = temp.rename(i)
        # 找到最高分、最低分、平均分
        s = pd.Series([df_temp['总分'].max(), df_temp['总分'].min(), df_temp['总分'].mean()],
                      index=['最高分', '最低分', '平均分'])
        # 用频道名称重新命名序列名
        s = s.rename(i)
        # 合并到等级序列中
        temp = temp.append(s)
        # print(temp)
        # 将各频道合并到一起
        fenxi.append(temp)
    # 生成pandas数据
    data = pd.DataFrame(fenxi)
    # 无数据填充为0
    data.fillna(0, inplace=True)
    # 增加频道各节目数
    temp = df.loc[:, '频道'].value_counts()
    # 将频道各节目数合并
    data.insert(0, '节目数量', temp)
    # 添加无数据列
    s = data.columns.to_list()
    dengji = ['节目数量', '优秀', '良好', '良', '及格', '不及格', '平均分', '最高分', '最低分']
    for i in dengji:
        if i in s:
            pass
        else:
            data[i] = 0
    data = data[dengji]
    data.insert(6, '优秀率', data[['节目数量', '优秀']].apply(lambda x: x['优秀'] / x['节目数量'], axis=1))
    data.insert(6, '达标率', data[['节目数量', '优秀', '良好', '良']].apply(
        lambda x: (x['优秀'] + x["良好"] + x['良']) / x['节目数量'], axis=1))
    # 数据类型
    dengji = ['节目数量', '优秀', '良好', '良', '及格', '不及格', '最高分', '最低分']
    data[dengji] = data[dengji].astype(np.int64)
    data['平均分'] = data['平均分'].round(2)
    data['优秀率%'] = data['优秀率'].apply(lambda x: format(x, '.2%'))
    data['达标率%'] = data['达标率'].apply(lambda x: format(x, '.2%'))
    data.reset_index(inplace=True)
    data = data.rename({'index': '频道'}, axis='columns')
    # print(data)
    # data.insert(0, '频道', pindao)
    # data.reset_index(drop=True,inplace=True)
    # print(data)

    # 表格数据写入Excel
    # 读取原数据
    souce = pd.read_excel('database.xlsx', sheet_name=None)
    new_sheet = '按频道分'
    if new_sheet in souce:
        souce.pop(new_sheet)
    with pd.ExcelWriter('database.xlsx', engine='xlsxwriter') as writer:
        for i in souce:
            souce[i].to_excel(writer, sheet_name=i, index=False)
        data.to_excel(writer, sheet_name=new_sheet, index=False)
        data[['频道', '达标率%', '优秀率%']].to_excel(writer, sheet_name='按频道分', startrow=data.shape[0] + 3, index=False)
        workbook = writer.book
        worksheet = writer.sheets[new_sheet]
        chart = workbook.add_chart({'type': 'column'})
        chart.add_series({
            'name': "=按频道分!$B$11",
            'categories': '=按频道分!$A$12:$C$18',
            'values': '=按频道分!$B$12:$B$18',
        })
        chart.add_series({
            'name': "=按频道分!$C$11",
            'categories': '=按频道分!$A$12:$C$18',
            'values': '=按频道分!$C$12:$C$18',
        })
        chart.set_title({'name': '各频道达标率、优秀率'})
        # chart.set_x_axis({'name': 'Test number'})
        # chart.set_y_axis({'name': 'Sample length (mm)'})
        chart.set_style(10)
        chart.height = 600
        chart.width = 960
        worksheet.insert_chart('D2', chart, {'x_offset': 25, 'y_offset': 10})

        data[['频道', '最高分', '最低分', '平均分']].to_excel(writer, sheet_name='按频道分',
                                                   startrow=(data.shape[0] + 3) * 2, index=False)
        chart = workbook.add_chart({'type': 'line'})
        chart.add_series({
            'name': "=按频道分!$B$21",
            'categories': '=按频道分!$A$22:$C$28',
            'values': '=按频道分!$B$22:$B$28',
        })
        chart.add_series({
            'name': "=按频道分!$C$21",
            'categories': '=按频道分!$A$22:$C$28',
            'values': '=按频道分!$C$22:$C$28',
        })
        chart.add_series({
            'name': "=按频道分!$D$21",
            'categories': '=按频道分!$A$22:$C$28',
            'values': '=按频道分!$D$22:$D$28',
        })
        chart.set_title({'name': '各频道分数比较'})
        # chart.set_x_axis({'name': 'Test number'})
        # chart.set_y_axis({'name': 'Sample length (mm)'})
        chart.set_style(10)
        chart.height = 600
        chart.width = 960
        worksheet.insert_chart('D38', chart, {'x_offset': 25, 'y_offset': 10})
    # 表格数据写入报告docx
    dengji = ['频道', '节目数量', '优秀', '良好', '良', '及格', '不及格', '达标率%', '优秀率%', '平均分']
    df = data[dengji]
    print(df)
    document = Document(file)
    # 将表格插入指定位置
    for p in document.paragraphs:
        if re.match("^Heading \d+$", p.style.name):
            if p.text == '二、综合得分等级及占比':
                print(p.text)
                table = document.add_table(rows=df.shape[0] + 2, cols=df.shape[1], style='Table Grid')
                move_table_after(table, p)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # 表头
    # 合并表头单元格
    for i in [0, 1, 7, 8, 9]:
        table.cell(0, i).merge(table.cell(1, i))
    cell1 = table.cell(0, 2).merge(table.cell(0, 4))
    cell1.text = '技术质量达标'
    cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 设置字体大小
    for run in cell1.paragraphs[0].runs:
        font = run.font
        font.size = Pt(10)
    cell1 = table.cell(0, 5).merge(table.cell(0, 6))
    cell1.text = '技术质量不达标'
    cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 设置表头底色
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value='#8DB4E2'))
    cell1._tc.get_or_add_tcPr().append(shading_elm_1)
    # 设置字体大小
    for run in cell1.paragraphs[0].runs:
        font = run.font
        font.size = Pt(10)
    # 各列宽度
    table_width = {'频道': 10.9, '节目数量': 14, '优秀': 14, '良好': 14, '良': 14,
                   '及格': 16, '不及格': 16, '达标率%': 17, '优秀率%': 16, '平均分': 16}
    # 取得各列名称
    columns_name = df.columns.to_list()
    for i in columns_name:
        # 设置表格列宽
        table.columns[columns_name.index(i)].width = Mm(table_width[i])
        cell1 = table.cell(1, columns_name.index(i))
        cell1.text = i
        cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 设置字体大小
        for run in cell1.paragraphs[0].runs:
            font = run.font
            font.size = Pt(10)

    for index, row in df.iterrows():
        for i in range(len(row)):
            cell1 = table.cell(index + 2, i)
            cell1.text = str(row[columns_name[i]])
            cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # 设置字体大小
            for run in cell1.paragraphs[0].runs:
                font = run.font
                font.size = Pt(10)
    '''
    # make picture
    dengji = ['达标率', '优秀率']
    df_lv = data[dengji]
    # df_lv[u'线损率'] = df_lv[u'线损率'].str.strip('%').astype(float) / 100
    df_lv.index = pindao
    print(df_lv.info())

    x_names = ['a','b','c']
    y_values = [1,2,3]
    plt_bar(x_names, y_values, "柱状图.png")
    plt_plot(x_names, y_values, "折线图.png")
    plt_scatter(x_names, y_values, "散点图.png")

    labels = 'Frogs', 'Hogs', 'Dogs', 'Logs'
    sizes = [15, 30, 45, 10]
    plt_pie(labels, sizes, "饼状图.png")

    document.add_picture('柱状图.png', width=Inches(6.25))
    document.add_picture('折线图.png', width=Inches(6.25))
    document.add_picture('饼状图.png', width=Inches(6.25))
    document.add_picture('散点图.png', width=Inches(6.25))


    #设置表格宋体大小
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(10)
    '''
    document.save(file)


def fenxi_youxiu(file='demo.docx'):
    df = pd.read_excel('database.xlsx')
    df = df[['节目名称', '播出时间', '总分']]
    df['播出时间'] = pd.to_datetime(df['播出时间'])
    df['播出时间'] = df['播出时间'].apply(lambda x: x.strftime('%Y年%m月%d日'))
    # 按总分排序
    df = df.sort_values(by='总分', ascending=False)
    df = df[df['总分'] >= 90]
    df.reset_index(drop=True, inplace=True)

    document = Document(file)
    # 将表格插入指定位置
    for p in document.paragraphs:
        if re.match("^Heading \d+$", p.style.name):
            if p.text == '1、优秀节目':
                print(p.text)
                # 插入表格
                table = document.add_table(rows=df.shape[0] + 1, cols=df.shape[1], style='Table Grid')
                # 移动表格到指定位置
                move_table_after(table, p)
    # 设置表格居中
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # 各列宽度
    table_width = {'排名': 7.4, '序号': 7.4, '节目名称': 50, '频道': 10.9, '播出时间': 33.2, '录制地点': 20.6,
                   '制作方式': 18, '制片人': 14.4, '主观': 7.4, '客观': 9.1, '总分': 11.3, '等级': 10.9}
    # 取得各列名称
    columns_name = df.columns.to_list()
    for i in columns_name:
        # 设置表格列宽
        table.columns[columns_name.index(i)].width = Mm(table_width[i])
        # 取得表格单元格
        cell1 = table.cell(0, columns_name.index(i))
        # 写入列名称
        cell1.text = i
        # 设置居中
        cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 设置垂直居中
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 设置表头底色
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value='#8DB4E2'))
        cell1._tc.get_or_add_tcPr().append(shading_elm_1)
        # 设置字体大小
        for run in cell1.paragraphs[0].runs:
            font = run.font
            font.size = Pt(10)
    table.rows[0].height = Mm(7.5)  # 表头行高
    for index, row in df.iterrows():
        table.rows[index + 1].height = Mm(7.5)  # 数据行高
        for i in range(len(row)):
            cell1 = table.cell(index + 1, i)
            cell1.text = str(row[columns_name[i]])
            cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # 设置字体大小
            for run in cell1.paragraphs[0].runs:
                font = run.font
                font.size = Pt(10)
    '''
    #设置表格宋体大小
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(10)
    '''
    document.save(file)


def fenxi_dabiao(file='demo.docx'):
    df = pd.read_excel('database.xlsx')
    df = df[['节目名称', '播出时间', '总分']]
    df['播出时间'] = pd.to_datetime(df['播出时间'])
    df['播出时间'] = df['播出时间'].apply(lambda x: x.strftime('%Y年%m月%d日'))
    # 按总分排序
    df = df.sort_values(by='总分', ascending=False)
    df = df[df['总分'] < 80]
    df.reset_index(drop=True, inplace=True)

    document = Document(file)
    # 将表格插入指定位置
    for p in document.paragraphs:
        if re.match("^Heading \d+$", p.style.name):
            if p.text == '2、不达标节目':
                print(p.text)
                # 插入表格
                table = document.add_table(rows=df.shape[0] + 1, cols=df.shape[1], style='Table Grid')
                # 移动表格到指定位置
                move_table_after(table, p)
    # 设置表格居中
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # 各列宽度
    table_width = {'排名': 7.4, '序号': 7.4, '节目名称': 50, '频道': 10.9, '播出时间': 33.2, '录制地点': 20.6,
                   '制作方式': 18, '制片人': 14.4, '主观': 7.4, '客观': 9.1, '总分': 11.3, '等级': 10.9}
    # 取得各列名称
    columns_name = df.columns.to_list()
    for i in columns_name:
        # 设置表格列宽
        table.columns[columns_name.index(i)].width = Mm(table_width[i])
        # 取得表格单元格
        cell1 = table.cell(0, columns_name.index(i))
        # 写入列名称
        cell1.text = i
        # 设置居中
        cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 设置垂直居中
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 设置表头底色
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value='#8DB4E2'))
        cell1._tc.get_or_add_tcPr().append(shading_elm_1)
        # 设置字体大小
        for run in cell1.paragraphs[0].runs:
            font = run.font
            font.size = Pt(10)
    table.rows[0].height = Mm(7.5)  # 表头行高
    for index, row in df.iterrows():
        table.rows[index + 1].height = Mm(7.5)  # 数据行高
        for i in range(len(row)):
            cell1 = table.cell(index + 1, i)
            cell1.text = str(row[columns_name[i]])
            cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # 设置字体大小
            for run in cell1.paragraphs[0].runs:
                font = run.font
                font.size = Pt(10)
    '''
    #设置表格宋体大小
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(10)
    '''
    document.save(file)


def zonghe_fen(file='demo.docx'):
    df = pd.read_excel('database.xlsx')
    # 按总分排序
    df = df.sort_values(by=['总分', '主观', '客观', '序号'], ascending=[False, False, False, True])
    df.reset_index(drop=True, inplace=True)
    df['排名'] = df.apply(lambda x: x.index + 1)

    df = df[['排名', '序号', '节目名称', '频道', '播出时间', '录制地点',
             '制作方式', '制片人', '主观', '客观', '总分', '等级']]
    df['主观'] = df['主观'].round(0).astype(np.int64)
    df['客观'] = df['客观'].astype(np.int64)
    df['总分'] = df['总分'].astype(np.int64)
    # 旧排序
    # df = df.sort_values(by='总分', ascending=False)
    # df.reset_index(drop=True, inplace=True)

    document = Document(file)
    # 将表格插入指定位置
    for p in document.paragraphs:
        if re.match("^Heading \d+$", p.style.name):
            if p.text == '二、综合得分排序':
                print(p.text)
                # 插入表格
                table = document.add_table(rows=df.shape[0] + 1, cols=df.shape[1], style='Table Grid')
                # 移动表格到指定位置
                move_table_after(table, p)
    # 设置表格居中
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # 各列宽度
    table_width = {'排名': 8.6, '序号': 8.6, '节目名称': 28, '频道': 11.2, '播出时间': 22.3, '录制地点': 20.9,
                   '制作方式': 18.2, '制片人': 19.3, '主观': 8.1, '客观': 10.1, '总分': 8, '等级': 7}
    # 取得各列名称
    columns_name = df.columns.to_list()
    for i in columns_name:
        # 设置表格列宽
        table.columns[columns_name.index(i)].width = Mm(table_width[i])
        # 取得表格单元格
        cell1 = table.cell(0, columns_name.index(i))
        # 写入列名称
        cell1.text = i
        # 设置居中
        cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 设置垂直居中
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 设置字体大小
        for run in cell1.paragraphs[0].runs:
            font = run.font
            font.size = Pt(10)
        # 设置表头底色
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value='#8DB4E2'))
        cell1._tc.get_or_add_tcPr().append(shading_elm_1)
    table.rows[0].height = Mm(15)  # 表头行高
    for index, row in df.iterrows():
        table.rows[index + 1].height = Mm(7.5)  # 数据行高
        for i in range(len(row)):
            cell1 = table.cell(index + 1, i)
            cell1.text = str(row[columns_name[i]])
            cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # 设置字体大小
            for run in cell1.paragraphs[0].runs:
                font = run.font
                font.size = Pt(10)
    # 合并频道单元格
    temp = df.loc[:, '等级'].value_counts()
    temp1 = temp.index.tolist()
    temp2 = ['优秀', '良好', '良', '及格', '不及格']
    dengji = []
    for i in temp2:
        if i in temp1:
            dengji.append(i)
    j = 1
    for i in dengji:
        cell1 = table.cell(j, 11).merge(table.cell(j + temp[i] - 1, 11))
        cell1.text = i
        cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        j += temp[i]
        # 设置字体大小
        for run in cell1.paragraphs[0].runs:
            font = run.font
            font.size = Pt(10)
    '''
    #设置表格宋体大小
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(10)
    '''
    document.save(file)


def rank_pindao(file='demo.docx'):
    df = pd.read_excel('database.xlsx')
    # 按总分排序
    df = df.sort_values(by=['总分', '主观', '客观', '序号'], ascending=[False, False, False, True])
    df.reset_index(drop=True, inplace=True)

    df = df[['频道', '节目名称', '播出时间', '录制地点',
             '制作方式', '制片人', '主观', '客观', '总分', '等级']]
    df['主观'] = df['主观'].round(0).astype(np.int64)
    df['客观'] = df['客观'].astype(np.int64)
    df['总分'] = df['总分'].astype(np.int64)

    pindao = ['卫视', '经济', '都市', '影视', '少儿', '公共', '农民']
    result = {}
    for i in pindao:
        # 筛选出频道数据
        df_temp = df[df['频道'] == i].copy()
        # 按总分排序  旧
        # df_temp = df_temp.sort_values(by='总分', ascending=False)
        # 重新生成行索引
        # df_temp.reset_index(drop=True, inplace=True)
        # 插入 排名 列
        df_temp.insert(1, '排名', df_temp['总分'].rank(ascending=False, method='first', ))
        # 排名列改为int32
        df_temp['排名'] = df_temp['排名'].astype(np.int32)
        # 保存到字典中
        result[i] = df_temp
    # 合并数据
    df = pd.concat(result)
    df.reset_index(drop=True, inplace=True)

    document = Document(file)
    # 将表格插入指定位置
    for p in document.paragraphs:
        if re.match("^Heading \d+$", p.style.name):
            if p.text == '一、按频道分析':
                print(p.text)
                # 插入表格
                table = document.add_table(rows=df.shape[0] + 1, cols=df.shape[1], style='Table Grid')
                # 移动表格到指定位置
                move_table_after(table, p)
    # 设置表格居中
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # 各列宽度
    table_width = {'排名': 9.1, '序号': 9.1, '节目名称': 29.6, '频道': 9.1, '播出时间': 23.6, '录制地点': 22.1,
                   '制作方式': 19.2, '制片人': 20.4, '主观': 8.6, '客观': 10.7, '总分': 8.4, '等级': 11.6}
    # 取得各列名称
    columns_name = df.columns.to_list()
    for i in columns_name:
        # 设置表格列宽
        table.columns[columns_name.index(i)].width = Mm(table_width[i])
        # 取得表格单元格
        cell1 = table.cell(0, columns_name.index(i))
        # 写入列名称
        cell1.text = i
        # 设置居中
        cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 设置垂直居中
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 设置字体大小
        for run in cell1.paragraphs[0].runs:
            font = run.font
            font.size = Pt(10)
        # 设置表头底色
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value='#8DB4E2'))
        cell1._tc.get_or_add_tcPr().append(shading_elm_1)
    # 写入数据
    table.rows[0].height = Mm(15)  # 表头行高
    for index, row in df.iterrows():
        table.rows[index + 1].height = Mm(7.5)  # 数据行高
        for i in range(len(row)):
            cell1 = table.cell(index + 1, i)
            cell1.text = str(row[columns_name[i]])
            cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # 设置字体大小
            for run in cell1.paragraphs[0].runs:
                font = run.font
                font.size = Pt(10)
    # 合并频道单元格
    temp = df.loc[:, '频道'].value_counts()
    j = 1
    for i in pindao:
        cell1 = table.cell(j, 0).merge(table.cell(j + temp[i] - 1, 0))
        cell1.text = i
        cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        j += temp[i]
        # 设置字体大小
        for run in cell1.paragraphs[0].runs:
            font = run.font
            font.size = Pt(10)
    '''
    #设置表格宋体大小
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(10)
    '''
    document.save(file)


def fenxi_pindao(file='demo.docx'):
    df = pd.read_excel('database.xlsx')
    df = df[['频道', '节目名称', '播出时间', '录制地点',
             '制作方式', '制片人', '主观', '客观', '总分', '等级']]
    df['主观'] = df['主观'].round(0)
    df['客观'] = df['客观'].astype(np.int64)
    df['总分'] = df['总分'].astype(np.int64)

    pindao = ['卫视', '经济', '都市', '影视', '少儿', '公共', '农民']
    fenxi = []
    for i in pindao:
        # 按频道筛选
        df_temp = df[df['频道'] == i]
        # 统计等级个数
        temp = df_temp.loc[:, '等级'].value_counts()
        # 用频道名称重新命名序列名
        temp = temp.rename(i)
        # 找到最高分、最低分、平均分
        s = pd.Series([df_temp['总分'].max(), df_temp['总分'].min(), df_temp['总分'].mean()],
                      index=['最高分', '最低分', '平均分'])
        # 用频道名称重新命名序列名
        s = s.rename(i)
        # 合并到等级序列中
        temp = temp.append(s)
        # print(temp)
        # 将各频道合并到一起
        fenxi.append(temp)
    # 生成pandas数据
    data = pd.DataFrame(fenxi)
    # 无数据填充为0
    data.fillna(0, inplace=True)
    # 增加频道各节目数
    temp = df.loc[:, '频道'].value_counts()
    # 将频道各节目数合并
    data.insert(0, '节目数量', temp)
    # 添加无数据列
    s = data.columns.to_list()
    dengji = ['节目数量', '优秀', '良好', '良', '及格', '不及格', '平均分', '最高分', '最低分']
    for i in dengji:
        if i in s:
            pass
        else:
            data[i] = 0
    data = data[dengji]
    data.insert(6, '优秀率', data[['节目数量', '优秀']].apply(lambda x: x['优秀'] / x['节目数量'], axis=1))
    data.insert(6, '达标率', data[['节目数量', '优秀', '良好', '良']].apply(
        lambda x: (x['优秀'] + x["良好"] + x['良']) / x['节目数量'], axis=1))
    # 数据类型
    dengji = ['节目数量', '优秀', '良好', '良', '及格', '不及格', '最高分', '最低分']
    data[dengji] = data[dengji].astype(np.int64)
    data['平均分'] = data['平均分'].round(2)
    data['达标率%'] = data['达标率'].apply(lambda x: format(x, '.2%'))
    data['优秀率%'] = data['优秀率'].apply(lambda x: format(x, '.2%'))
    data.reset_index(inplace=True)
    data = data.rename({'index': '频道'}, axis='columns')
    # print(data)
    # data.insert(0, '频道', pindao)
    # data.reset_index(drop=True,inplace=True)
    # print(data)

    # 表格数据写入Excel
    # 读取原数据
    souce = pd.read_excel('database.xlsx', sheet_name=None)
    new_sheet = '按频道分'
    if new_sheet in souce:
        souce.pop(new_sheet)
    with pd.ExcelWriter('database.xlsx', engine='xlsxwriter') as writer:
        for i in souce:
            souce[i].to_excel(writer, sheet_name=i, index=False)
        data.to_excel(writer, sheet_name=new_sheet, index=False)
        data[['频道', '达标率%', '优秀率%']].to_excel(writer, sheet_name='按频道分', startrow=data.shape[0] + 3, index=False)
        workbook = writer.book
        worksheet = writer.sheets[new_sheet]
        chart = workbook.add_chart({'type': 'column'})
        chart.add_series({
            'name': "=按频道分!$B$11",
            'categories': '=按频道分!$A$12:$C$18',
            'values': '=按频道分!$B$12:$B$18',
        })
        chart.add_series({
            'name': "=按频道分!$C$11",
            'categories': '=按频道分!$A$12:$C$18',
            'values': '=按频道分!$C$12:$C$18',
        })
        chart.set_title({'name': '各频道达标率、优秀率'})
        # chart.set_x_axis({'name': 'Test number'})
        # chart.set_y_axis({'name': 'Sample length (mm)'})
        chart.set_style(10)
        chart.height = 600
        chart.width = 960
        worksheet.insert_chart('D2', chart, {'x_offset': 25, 'y_offset': 10})

        data[['频道', '最高分', '最低分', '平均分']].to_excel(writer, sheet_name='按频道分',
                                                   startrow=(data.shape[0] + 3) * 2, index=False)
        chart = workbook.add_chart({'type': 'line'})
        chart.add_series({
            'name': "=按频道分!$B$21",
            'categories': '=按频道分!$A$22:$C$28',
            'values': '=按频道分!$B$22:$B$28',
        })
        chart.add_series({
            'name': "=按频道分!$C$21",
            'categories': '=按频道分!$A$22:$C$28',
            'values': '=按频道分!$C$22:$C$28',
        })
        chart.add_series({
            'name': "=按频道分!$D$21",
            'categories': '=按频道分!$A$22:$C$28',
            'values': '=按频道分!$D$22:$D$28',
        })
        chart.set_title({'name': '各频道分数比较'})
        # chart.set_x_axis({'name': 'Test number'})
        # chart.set_y_axis({'name': 'Sample length (mm)'})
        chart.set_style(10)
        chart.height = 600
        chart.width = 960
        worksheet.insert_chart('D38', chart, {'x_offset': 25, 'y_offset': 10})
    # 表格数据写入报告docx
    dengji = ['频道', '节目数量', '优秀', '良好', '良', '及格', '不及格', '达标率%', '优秀率%', '平均分']
    df = data[dengji]
    print(df)
    document = Document(file)
    # 将表格插入指定位置
    for p in document.paragraphs:
        if re.match("^Heading \d+$", p.style.name):
            if p.text == '频道分析表':
                print(p.text)
                table = document.add_table(rows=df.shape[0] + 2, cols=df.shape[1], style='Table Grid')
                move_table_after(table, p)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # 表头
    # 合并表头单元格
    for i in [0, 1, 7, 8, 9]:
        table.cell(0, i).merge(table.cell(1, i))
    cell1 = table.cell(0, 2).merge(table.cell(0, 4))
    cell1.text = '技术质量达标'
    cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 设置表头底色
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value='#D6E3BC'))
    cell1._tc.get_or_add_tcPr().append(shading_elm_1)
    # 设置字体大小
    for run in cell1.paragraphs[0].runs:
        font = run.font
        font.size = Pt(10)
    cell1 = table.cell(0, 5).merge(table.cell(0, 6))
    cell1.text = '技术质量不达标'
    cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 设置表头底色
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value='#E5B8B7'))
    cell1._tc.get_or_add_tcPr().append(shading_elm_1)
    # 设置字体大小
    for run in cell1.paragraphs[0].runs:
        font = run.font
        font.size = Pt(10)
    # 各列宽度
    table_width = {'频道': 10.9, '节目数量': 14, '优秀': 14, '良好': 14, '良': 14,
                   '及格': 16, '不及格': 16, '达标率%': 17, '优秀率%': 17, '平均分': 16}
    table_colors = ['#8DB3E2', '#8DB3E2', '#3AA315', '#9BBB59', '#943634', '#C0504D', '#D8D8D8',
                    '#C6D9F1', '#C6D9F1', '#C6D9F1']
    # 取得各列名称
    columns_name = df.columns.to_list()
    for i in columns_name:
        # 设置表格列宽
        table.columns[columns_name.index(i)].width = Mm(table_width[i])
        cell1 = table.cell(1, columns_name.index(i))
        cell1.text = i
        cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 设置标题颜色
        table_color = table_colors[columns_name.index(i)]
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value=table_color))
        cell1._tc.get_or_add_tcPr().append(shading_elm_1)
        # 设置字体大小
        for run in cell1.paragraphs[0].runs:
            font = run.font
            font.size = Pt(10)

    for index, row in df.iterrows():
        for i in range(len(row)):
            cell1 = table.cell(index + 2, i)
            cell1.text = str(row[columns_name[i]])
            cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # 设置字体大小
            for run in cell1.paragraphs[0].runs:
                font = run.font
                font.size = Pt(10)
    '''
    # make picture
    dengji = ['达标率', '优秀率']
    df_lv = data[dengji]
    # df_lv[u'线损率'] = df_lv[u'线损率'].str.strip('%').astype(float) / 100
    df_lv.index = pindao
    print(df_lv.info())

    x_names = ['a','b','c']
    y_values = [1,2,3]
    plt_bar(x_names, y_values, "柱状图.png")
    plt_plot(x_names, y_values, "折线图.png")
    plt_scatter(x_names, y_values, "散点图.png")

    labels = 'Frogs', 'Hogs', 'Dogs', 'Logs'
    sizes = [15, 30, 45, 10]
    plt_pie(labels, sizes, "饼状图.png")

    document.add_picture('柱状图.png', width=Inches(6.25))
    document.add_picture('折线图.png', width=Inches(6.25))
    document.add_picture('饼状图.png', width=Inches(6.25))
    document.add_picture('散点图.png', width=Inches(6.25))

    
    #设置表格宋体大小
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(10)
    '''
    document.save(file)


def rank_didian(file='demo.docx'):
    df = pd.read_excel('database.xlsx')
    # 按总分排序
    df = df.sort_values(by=['总分', '主观', '客观', '序号'], ascending=[False, False, False, True])
    df.reset_index(drop=True, inplace=True)

    didian = ['引进节目', '外景录制', '台外录制', '台内录制']
    result = {}
    for i in didian:
        # 筛选出频道数据
        df_temp = df[df['地点'] == i].copy()
        # 按总分排序
        if i == '台内录制':
            df_temp['演播室'] = df_temp['录制地点'].apply(lambda x: x[:-3])
            df_temp['演播室'] = df_temp['演播室'].astype(np.int64)
            df_temp = df_temp.sort_values(by=['演播室', '总分', '主观', '客观', '序号'],
                                          ascending=[False, False, False, False, True])
            df_temp.drop(['演播室'], axis=1, inplace=True)
        # else:
        #     df_temp = df_temp.sort_values(by='总分', ascending=False)
        # 重新生成行索引
        # df_temp.reset_index(drop=True, inplace=True)
        # 插入 排名 列
        # df_temp.insert(1, '排名', df_temp['总分'].rank(ascending=False, method='first',))
        # 排名列改为int32
        # df_temp['排名'] = df_temp['排名'].astype(np.int32)
        # 保存到字典中
        result[i] = df_temp

    # print(result)
    # 合并数据
    df = pd.concat(result)
    df.reset_index(drop=True, inplace=True)
    # 选择要展示字段
    df = df[['地点', '录制地点', '节目名称', '频道', '总分', '等级']]
    # df['主观'] = df['主观'].round(0).astype(np.int64)
    # df['客观'] = df['客观'].astype(np.int64)
    df['总分'] = df['总分'].astype(np.int64)

    document = Document(file)
    # 将表格插入指定位置
    for p in document.paragraphs:
        if re.match("^Heading \d+$", p.style.name):
            if p.text == '二、按录制地点分析':
                print(p.text)
                # 插入表格
                table = document.add_table(rows=df.shape[0] + 1, cols=df.shape[1], style='Table Grid')
                # 移动表格到指定位置
                move_table_after(table, p)
    # 设置表格居中
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # 各列宽度
    table_width = {'地点': 20.6, '频道': 7.4, '节目名称': 60, '频道': 10.9, '播出时间': 21.5, '录制地点': 23.5,
                   '制作方式': 18, '制片人': 14.4, '主观': 7.4, '客观': 9.1, '总分': 10.9, '等级': 10.9}
    # 取得各列名称
    columns_name = df.columns.to_list()
    for i in columns_name:
        # 设置表格列宽
        table.columns[columns_name.index(i)].width = Mm(table_width[i])
        # 取得表格单元格
        cell1 = table.cell(0, columns_name.index(i))
        # 写入列名称
        cell1.text = i
        # 设置居中
        cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 设置垂直居中
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 设置字体大小
        for run in cell1.paragraphs[0].runs:
            font = run.font
            font.size = Pt(10)
        # 设置表头底色
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value='#8DB4E2'))
        cell1._tc.get_or_add_tcPr().append(shading_elm_1)
    # 写入数据
    table.rows[0].height = Mm(7.5)  # 表头行高
    for index, row in df.iterrows():
        table.rows[index + 1].height = Mm(7.5)  # 数据行高
        for i in range(len(row)):
            cell1 = table.cell(index + 1, i)
            cell1.text = str(row[columns_name[i]])
            cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # 设置字体大小
            for run in cell1.paragraphs[0].runs:
                font = run.font
                font.size = Pt(10)
    # 合并地点单元格
    temp = df.loc[:, '地点'].value_counts()
    j = 1
    for i in didian:
        cell1 = table.cell(j, 0).merge(table.cell(j + temp[i] - 1, 0))
        cell1.text = i
        cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        j += temp[i]
        # 设置字体大小
        for run in cell1.paragraphs[0].runs:
            font = run.font
            font.size = Pt(10)
    # 合并演播室单元格
    temp = df_temp['录制地点'].value_counts()
    j = len(df) - len(df_temp) + 1
    tai_inter = ['800演播室', '400演播室', '300演播室', '260演播室', '120演播室', '110演播室', '70演播室', ]
    for i in tai_inter:
        cell1 = table.cell(j, 1).merge(table.cell(j + temp[i] - 1, 1))
        cell1.text = i
        cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        j += temp[i]
        # 设置字体大小
        for run in cell1.paragraphs[0].runs:
            font = run.font
            font.size = Pt(10)
    '''
    #设置表格宋体大小
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(10)
    '''
    document.save(file)


def fenxi_didian(file='demo.docx'):
    df = pd.read_excel('database.xlsx')
    df = df[['地点', '录制地点', '节目名称', '频道', '总分', '等级']]
    # df['主观'] = df['主观'].round(0).astype(np.int64)
    # df['客观'] = df['客观'].astype(np.int64)
    df['总分'] = df['总分'].astype(np.int64)

    fenxi = []
    didian = ['引进节目', '外景录制', '台外录制', '台内录制']
    for i in didian:
        # 按频道筛选
        df_temp = df[df['地点'] == i]
        # 统计等级个数
        temp = df_temp.loc[:, '等级'].value_counts()
        # 用频道名称重新命名序列名
        temp = temp.rename(i)
        # 找到最高分、最低分、平均分
        s = pd.Series([df_temp['总分'].max(), df_temp['总分'].min(), df_temp['总分'].mean()],
                      index=['最高分', '最低分', '平均分'])
        # 用频道名称重新命名序列名
        s = s.rename(i)
        # 合并到等级序列中
        temp = temp.append(s)
        # print(temp)
        # 将各频道合并到一起
        fenxi.append(temp)
    tai_inter = ['800演播室', '400演播室', '300演播室', '260演播室', '120演播室', '110演播室', '70演播室', ]
    for i in tai_inter:
        # 按频道筛选
        df_temp = df[df['录制地点'] == i]
        # 统计等级个数
        temp = df_temp.loc[:, '等级'].value_counts()
        # 用频道名称重新命名序列名
        temp = temp.rename(i)
        # 找到最高分、最低分、平均分
        s = pd.Series([df_temp['总分'].max(), df_temp['总分'].min(), df_temp['总分'].mean()],
                      index=['最高分', '最低分', '平均分'])
        # 用频道名称重新命名序列名
        s = s.rename(i)
        # 合并到等级序列中
        temp = temp.append(s)
        # print(temp)
        # 将各频道合并到一起
        fenxi.append(temp)
    # 生成pandas数据
    data = pd.DataFrame(fenxi)
    # 无数据填充为0
    data.fillna(0, inplace=True)
    # 增加频道各节目数
    temp = df.loc[:, '地点'].value_counts()
    temp1 = df[df['地点'] == '台内录制'].loc[:, '录制地点'].value_counts()
    temp = temp.append(temp1)
    # 将频道各节目数合并
    data.insert(0, '节目数量', temp)
    # 添加无数据列
    s = data.columns.to_list()
    dengji = ['节目数量', '优秀', '良好', '良', '及格', '不及格', '平均分', '最高分', '最低分']
    for i in dengji:
        if i in s:
            pass
        else:
            data[i] = 0
    data = data[dengji]
    data.insert(6, '优秀率', data[['节目数量', '优秀']].apply(lambda x: x['优秀'] / x['节目数量'], axis=1))
    data.insert(6, '达标率', data[['节目数量', '优秀', '良好', '良']].apply(
        lambda x: (x['优秀'] + x["良好"] + x['良']) / x['节目数量'], axis=1))
    # 数据类型
    dengji = ['节目数量', '优秀', '良好', '良', '及格', '不及格', '最高分', '最低分']
    data[dengji] = data[dengji].astype(np.int64)
    data['平均分'] = data['平均分'].round(2)
    data['达标率%'] = data['达标率'].apply(lambda x: format(x, '.2%'))
    data['优秀率%'] = data['优秀率'].apply(lambda x: format(x, '.2%'))
    data.reset_index(inplace=True)
    data = data.rename({'index': '地点'}, axis='columns')

    # print(data)
    # data.insert(0, '频道', pindao)
    # data.reset_index(drop=True,inplace=True)
    # print(data)

    # 表格数据写入Excel
    # 读取原数据
    souce = pd.read_excel('database.xlsx', sheet_name=None)
    new_sheet = '按地点分'
    if new_sheet in souce:
        souce.pop(new_sheet)
    with pd.ExcelWriter('database.xlsx', engine='xlsxwriter') as writer:
        for i in souce:
            souce[i].to_excel(writer, sheet_name=i, index=False)
        data.to_excel(writer, sheet_name=new_sheet, index=False)
        # 图表数据
        tu1 = data[['地点', '达标率%', '优秀率%']]
        tu1 = tu1[tu1['地点'].isin(didian)]
        tu1.to_excel(writer, sheet_name='按地点分', startrow=14, index=False)
        tu2 = data[['地点', '达标率%', '优秀率%']]
        tu2 = tu2[tu2['地点'].isin(tai_inter)]
        tu2.to_excel(writer, sheet_name='按地点分', startrow=19, index=False)

        workbook = writer.book
        worksheet = writer.sheets[new_sheet]
        chart = workbook.add_chart({'type': 'column'})
        chart.add_series({
            'name': "=按地点分!$B$15",
            'categories': '=按地点分!$A$16:$C$19',
            'values': '=按地点分!$B$16:$B$19',
        })
        chart.add_series({
            'name': "=按地点分!$C$15",
            'categories': '=按地点分!$A$16:$C$19',
            'values': '=按地点分!$C$16:$C$19',
        })
        chart.set_title({'name': '各录制地点达标率、优秀率'})
        # chart.set_x_axis({'name': 'Test number'})
        # chart.set_y_axis({'name': 'Sample length (mm)'})
        chart.set_style(10)
        chart.height = 600
        chart.width = 960
        worksheet.insert_chart('D2', chart, {'x_offset': 25, 'y_offset': 10})

    # 表格数据写入报告docx
    dengji = ['地点', '节目数量', '优秀', '良好', '良', '及格', '不及格', '达标率%', '优秀率%', ]
    df = data[dengji]
    print(df)
    document = Document(file)
    # 将表格插入指定位置
    for p in document.paragraphs:
        if re.match("^Heading \d+$", p.style.name):
            if p.text == '录制地点分析表':
                print(p.text)
                table = document.add_table(rows=df.shape[0] + 2, cols=df.shape[1], style='Table Grid')
                move_table_after(table, p)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # 表头
    # 合并表头单元格
    for i in [0, 1, 7, 8]:
        table.cell(0, i).merge(table.cell(1, i))
    cell1 = table.cell(0, 2).merge(table.cell(0, 4))
    cell1.text = '技术质量达标'
    cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 设置表头底色
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value='#D6E3BC'))
    cell1._tc.get_or_add_tcPr().append(shading_elm_1)
    # 设置字体大小
    for run in cell1.paragraphs[0].runs:
        font = run.font
        font.size = Pt(10)
    cell1 = table.cell(0, 5).merge(table.cell(0, 6))
    cell1.text = '技术质量不达标'
    cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 设置表头底色
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value='#E5B8B7'))
    cell1._tc.get_or_add_tcPr().append(shading_elm_1)
    # 设置字体大小
    for run in cell1.paragraphs[0].runs:
        font = run.font
        font.size = Pt(10)
    # 各列宽度
    table_width = {'地点': 20.8, '节目数量': 14, '优秀': 14, '良好': 14, '良': 14,
                   '及格': 16, '不及格': 16, '达标率%': 19, '优秀率%': 19, '平均分': 16}
    table_colors = ['#8DB3E2', '#8DB3E2', '#3AA315', '#9BBB59', '#943634', '#C0504D', '#D8D8D8',
                    '#C6D9F1', '#C6D9F1', '#C6D9F1']
    # 取得各列名称
    columns_name = df.columns.to_list()
    for i in columns_name:
        # 设置表格列宽
        table.columns[columns_name.index(i)].width = Mm(table_width[i])
        cell1 = table.cell(1, columns_name.index(i))
        cell1.text = i
        cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 设置标题颜色
        table_color = table_colors[columns_name.index(i)]
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value=table_color))
        cell1._tc.get_or_add_tcPr().append(shading_elm_1)
        # 设置字体大小
        for run in cell1.paragraphs[0].runs:
            font = run.font
            font.size = Pt(10)

    for index, row in df.iterrows():
        for i in range(len(row)):
            cell1 = table.cell(index + 2, i)
            cell1.text = str(row[columns_name[i]])
            cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # 设置字体大小
            for run in cell1.paragraphs[0].runs:
                font = run.font
                font.size = Pt(10)
    '''
    # make picture
    dengji = ['达标率', '优秀率']
    df_lv = data[dengji]
    # df_lv[u'线损率'] = df_lv[u'线损率'].str.strip('%').astype(float) / 100
    df_lv.index = pindao
    print(df_lv.info())

    x_names = ['a','b','c']
    y_values = [1,2,3]
    plt_bar(x_names, y_values, "柱状图.png")
    plt_plot(x_names, y_values, "折线图.png")
    plt_scatter(x_names, y_values, "散点图.png")

    labels = 'Frogs', 'Hogs', 'Dogs', 'Logs'
    sizes = [15, 30, 45, 10]
    plt_pie(labels, sizes, "饼状图.png")

    document.add_picture('柱状图.png', width=Inches(6.25))
    document.add_picture('折线图.png', width=Inches(6.25))
    document.add_picture('饼状图.png', width=Inches(6.25))
    document.add_picture('散点图.png', width=Inches(6.25))


    #设置表格宋体大小
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(10)
    '''
    document.save(file)


def rank_fangshi(file='demo.docx'):
    df = pd.read_excel('database.xlsx')
    # 按总分排序
    df = df.sort_values(by=['总分', '主观', '客观', '序号'], ascending=[False, False, False, True])
    df.reset_index(drop=True, inplace=True)

    # df['主观'] = df['主观'].round(0).astype(np.int64)
    # df['客观'] = df['客观'].astype(np.int64)
    df['总分'] = df['总分'].astype(np.int64)
    # group方式，无法控制前后顺序
    # tt = df.groupby(['方式', '制作方式']).apply(lambda x: x.sort_values('总分', ascending=False))
    fangshi = ['台外制作', '台内制作']
    # 制作方式分类
    tai_outer = ['北京制作', '台外制作', '联合制作']
    tai_inter = ['包装制作', '高清自制', '直播', '120自制', '影视自制', '少儿自制', '农民自制', '广告自制', '录播']
    result = {}
    for i in fangshi:
        # 筛选出频道数据
        df_temp = df[df['方式'] == i].copy()
        # 按总分排序
        if i == '台内制作':
            df_temp['方式1'] = df_temp['制作方式'].apply((lambda x: tai_inter.index(x)))
        else:
            df_temp['方式1'] = df_temp['制作方式'].apply((lambda x: tai_outer.index(x)))
        df_temp = df_temp.sort_values(by=['方式1', '总分', '主观', '客观', '序号'],
                                      ascending=[True, False, False, False, True])
        df_temp.drop(['方式1'], axis=1, inplace=True)
        # 重新生成行索引
        # df_temp.reset_index(drop=True, inplace=True)
        # 插入 排名 列
        # df_temp.insert(1, '排名', df_temp['总分'].rank(ascending=False, method='first',))
        # 排名列改为int32
        # df_temp['排名'] = df_temp['排名'].astype(np.int32)
        # 保存到字典中
        result[i] = df_temp
    # 合并数据
    df = pd.concat(result)
    df.reset_index(drop=True, inplace=True)
    # 选择展示字段
    df = df[['方式', '制作方式', '节目名称', '频道', '总分', '等级']]
    document = Document(file)
    # 将表格插入指定位置
    for p in document.paragraphs:
        if re.match("^Heading \d+$", p.style.name):
            if p.text == '三、按制作方式分析':
                print(p.text)
                # 插入表格
                table = document.add_table(rows=df.shape[0] + 1, cols=df.shape[1], style='Table Grid')
                # 移动表格到指定位置
                move_table_after(table, p)
    # 设置表格居中
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # 各列宽度
    table_width = {'方式': 20.6, '节目名称': 60, '频道': 10.9, '播出时间': 21.5, '录制地点': 23.5,
                   '制作方式': 20, '制片人': 14.4, '主观': 7.4, '客观': 9.1, '总分': 10.9, '等级': 10.9}
    # 取得各列名称
    columns_name = df.columns.to_list()
    for i in columns_name:
        # 设置表格列宽
        table.columns[columns_name.index(i)].width = Mm(table_width[i])
        # 取得表格单元格
        cell1 = table.cell(0, columns_name.index(i))
        # 写入列名称
        cell1.text = i
        # 设置居中
        cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 设置垂直居中
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 设置表头底色
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value='#8DB4E2'))
        cell1._tc.get_or_add_tcPr().append(shading_elm_1)
        # 设置字体大小
        for run in cell1.paragraphs[0].runs:
            font = run.font
            font.size = Pt(10)

    # 写入数据
    table.rows[0].height = Mm(7.5)  # 表头行高
    for index, row in df.iterrows():
        table.rows[index + 1].height = Mm(7.5)  # 数据行高
        for i in range(len(row)):
            cell1 = table.cell(index + 1, i)
            cell1.text = str(row[columns_name[i]])
            cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # 设置字体大小
            for run in cell1.paragraphs[0].runs:
                font = run.font
                font.size = Pt(10)
    # 合并方式单元格
    temp = df.loc[:, '方式'].value_counts()
    j = 1
    for i in fangshi:
        cell1 = table.cell(j, 0).merge(table.cell(j + temp[i] - 1, 0))
        cell1.text = i
        cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        j += temp[i]
        # 设置字体大小
        for run in cell1.paragraphs[0].runs:
            font = run.font
            font.size = Pt(10)
    # 合并制作方式单元格
    temp = df['制作方式'].value_counts()  # 制作方式行数统计
    # 合并制作方式list
    tai_outer.extend(tai_inter)
    # 去掉此报告中没有的制作方式
    tai = []
    for i in tai_outer:
        if i in temp.index.to_list():
            tai.append(i)

    j = 1  # 合并制作方式开始的单元格行数
    for i in tai:
        cell1 = table.cell(j, 1).merge(table.cell(j + temp[i] - 1, 1))
        cell1.text = i
        cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        j += temp[i]
        # 设置字体大小
        for run in cell1.paragraphs[0].runs:
            font = run.font
            font.size = Pt(10)
    '''
    #设置表格宋体大小
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(10)
    '''
    document.save(file)


def fenxi_fangshi(file='demo.docx'):
    df = pd.read_excel('database.xlsx')
    df = df[['方式', '制作方式', '节目名称', '频道', '总分', '等级']]
    # df['主观'] = df['主观'].round(0).astype(np.int64)
    # df['客观'] = df['客观'].astype(np.int64)
    df['总分'] = df['总分'].astype(np.int64)

    fenxi = []
    fangshi = ['台外制作', '台内制作']
    tai_outer = ['北京制作', '台外制作', '联合制作']
    tai_inter = ['包装制作', '高清自制', '直播', '120自制', '影视自制', '少儿自制', '农民自制', '广告自制', '录播']
    tai_outer.extend(tai_inter)
    for i in fangshi:
        # 按频道筛选
        df_temp = df[df['方式'] == i]
        # 统计等级个数
        temp = df_temp.loc[:, '等级'].value_counts()
        # 用频道名称重新命名序列名
        temp = temp.rename(i)
        # 找到最高分、最低分、平均分
        s = pd.Series([df_temp['总分'].max(), df_temp['总分'].min(), df_temp['总分'].mean()],
                      index=['最高分', '最低分', '平均分'])
        # 用频道名称重新命名序列名
        s = s.rename(i)
        # 合并到等级序列中
        temp = temp.append(s)
        # print(temp)
        # 将各频道合并到一起
        fenxi.append(temp)
        # 二级目录
        temp1 = df_temp.drop_duplicates(subset='制作方式', keep='first')
        temp2 = temp1['制作方式'].tolist()
        temp = []
        for k in tai_outer:
            if k in temp2:
                temp.append(k)

        for j in temp:
            # 按频道筛选
            df_temp = df[df['制作方式'] == j]
            # 统计等级个数
            temp = df_temp.loc[:, '等级'].value_counts()
            # 用频道名称重新命名序列名
            temp = temp.rename(j)
            # 找到最高分、最低分、平均分
            s = pd.Series([df_temp['总分'].max(), df_temp['总分'].min(), df_temp['总分'].mean()],
                          index=['最高分', '最低分', '平均分'])
            # 用频道名称重新命名序列名
            s = s.rename(j)
            # 合并到等级序列中
            temp = temp.append(s)
            # print(temp)
            # 将各频道合并到一起
            fenxi.append(temp)
    # 生成pandas数据
    data = pd.DataFrame(fenxi)
    # 无数据填充为0
    data.fillna(0, inplace=True)
    # 添加无数据列
    s = data.columns.to_list()
    dengji = ['节目数量', '优秀', '良好', '良', '及格', '不及格', '平均分', '最高分', '最低分']
    for i in dengji:
        if i in s:
            pass
        else:
            data[i] = 0
    # 增加节目数汇总
    data['节目数量'] = data[['优秀', '良好', '良', '及格', '不及格']].sum(axis=1)

    # data = data[dengji]
    data.insert(6, '优秀率', data[['节目数量', '优秀']].apply(lambda x: x['优秀'] / x['节目数量'], axis=1))
    data.insert(6, '达标率', data[['节目数量', '优秀', '良好', '良']].apply(
        lambda x: (x['优秀'] + x["良好"] + x['良']) / x['节目数量'], axis=1))
    # 数据类型
    dengji = ['节目数量', '优秀', '良好', '良', '及格', '不及格', '最高分', '最低分']
    data[dengji] = data[dengji].astype(np.int64)
    data['平均分'] = data['平均分'].round(2)
    data['达标率%'] = data['达标率'].apply(lambda x: format(x, '.2%'))
    data['优秀率%'] = data['优秀率'].apply(lambda x: format(x, '.2%'))
    data.reset_index(inplace=True)
    data = data.rename({'index': '方式'}, axis='columns')

    # data.insert(0, '频道', pindao)
    # data.reset_index(drop=True,inplace=True)
    # print(data)

    # 表格数据写入Excel
    # 读取原数据
    souce = pd.read_excel('database.xlsx', sheet_name=None)
    new_sheet = '按方式分'
    if new_sheet in souce:
        souce.pop(new_sheet)
    with pd.ExcelWriter('database.xlsx', engine='xlsxwriter') as writer:
        for i in souce:
            souce[i].to_excel(writer, sheet_name=i, index=False)
        data.to_excel(writer, sheet_name=new_sheet, index=False)
        # 图表数据
        tu1 = data[['方式', '达标率%', '优秀率%']]
        tu1 = tu1[tu1['方式'].isin(tai_outer)]
        tu1.to_excel(writer, sheet_name='按方式分', startrow=14, index=False)
        tu2 = data[['方式', '达标率%', '优秀率%']]
        tu2 = tu2[tu2['方式'].isin(tai_inter)]
        tu2.to_excel(writer, sheet_name='按方式分', startrow=19, index=False)
        # 插入图表
        workbook = writer.book
        worksheet = writer.sheets[new_sheet]
        chart = workbook.add_chart({'type': 'column'})
        chart.add_series({
            'name': "=按方式分!$B$11",
            'categories': '=按方式分!$A$12:$C$18',
            'values': '=按方式分!$B$12:$B$18',
        })
        chart.add_series({
            'name': "=按方式分!$C$11",
            'categories': '=按方式分!$A$12:$C$18',
            'values': '=按方式分!$C$12:$C$18',
        })
        chart.set_title({'name': '各制作方式达标率、优秀率'})
        # chart.set_x_axis({'name': 'Test number'})
        # chart.set_y_axis({'name': 'Sample length (mm)'})
        chart.set_style(10)
        chart.height = 600
        chart.width = 960
        worksheet.insert_chart('D2', chart, {'x_offset': 25, 'y_offset': 10})

    # 表格数据写入报告docx
    dengji = ['方式', '节目数量', '优秀', '良好', '良', '及格', '不及格', '达标率%', '优秀率%', '平均分']
    df = data[dengji]
    print(df)
    document = Document(file)
    # 将表格插入指定位置
    for p in document.paragraphs:
        if re.match("^Heading \d+$", p.style.name):
            if p.text == '制作方式分析表':
                print(p.text)
                table = document.add_table(rows=df.shape[0] + 2, cols=df.shape[1], style='Table Grid')
                move_table_after(table, p)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # 表头
    # 合并表头单元格
    for i in [0, 1, 7, 8, 9]:
        table.cell(0, i).merge(table.cell(1, i))
    cell1 = table.cell(0, 2).merge(table.cell(0, 4))
    cell1.text = '技术质量达标'
    cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 设置表头底色
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value='#D6E3BC'))
    cell1._tc.get_or_add_tcPr().append(shading_elm_1)
    # 设置字体大小
    for run in cell1.paragraphs[0].runs:
        font = run.font
        font.size = Pt(10)
    cell1 = table.cell(0, 5).merge(table.cell(0, 6))
    cell1.text = '技术质量不达标'
    cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 设置表头底色
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value='#E5B8B7'))
    cell1._tc.get_or_add_tcPr().append(shading_elm_1)
    # 设置字体大小
    for run in cell1.paragraphs[0].runs:
        font = run.font
        font.size = Pt(10)
    # 各列宽度
    table_width = {'方式': 20.8, '节目数量': 14, '优秀': 14, '良好': 14, '良': 14,
                   '及格': 16, '不及格': 16, '达标率%': 17, '优秀率%': 17, '平均分': 16}
    table_colors = ['#8DB3E2', '#8DB3E2', '#3AA315', '#9BBB59', '#943634', '#C0504D', '#D8D8D8',
                    '#C6D9F1', '#C6D9F1', '#C6D9F1']
    # 取得各列名称
    columns_name = df.columns.to_list()
    for i in columns_name:
        # 设置表格列宽
        table.columns[columns_name.index(i)].width = Mm(table_width[i])
        cell1 = table.cell(1, columns_name.index(i))
        cell1.text = i
        cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 设置标题颜色
        table_color = table_colors[columns_name.index(i)]
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value=table_color))
        cell1._tc.get_or_add_tcPr().append(shading_elm_1)
        # 设置字体大小
        for run in cell1.paragraphs[0].runs:
            font = run.font
            font.size = Pt(10)

    for index, row in df.iterrows():
        for i in range(len(row)):
            cell1 = table.cell(index + 2, i)
            cell1.text = str(row[columns_name[i]])
            cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # 设置字体大小
            for run in cell1.paragraphs[0].runs:
                font = run.font
                font.size = Pt(10)
    '''
    # make picture
    dengji = ['达标率', '优秀率']
    df_lv = data[dengji]
    # df_lv[u'线损率'] = df_lv[u'线损率'].str.strip('%').astype(float) / 100
    df_lv.index = pindao
    print(df_lv.info())

    x_names = ['a','b','c']
    y_values = [1,2,3]
    plt_bar(x_names, y_values, "柱状图.png")
    plt_plot(x_names, y_values, "折线图.png")
    plt_scatter(x_names, y_values, "散点图.png")

    labels = 'Frogs', 'Hogs', 'Dogs', 'Logs'
    sizes = [15, 30, 45, 10]
    plt_pie(labels, sizes, "饼状图.png")

    document.add_picture('柱状图.png', width=Inches(6.25))
    document.add_picture('折线图.png', width=Inches(6.25))
    document.add_picture('饼状图.png', width=Inches(6.25))
    document.add_picture('散点图.png', width=Inches(6.25))


    #设置表格宋体大小
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(10)
    '''
    document.save(file)


def Experts_zongping(file='demo.docx'):
    df = pd.read_excel('database.xlsx')
    # 按总分排序
    df = df.sort_values(by=['总分', '主观', '客观', '序号'], ascending=[False, False, False, True])
    df.reset_index(drop=True, inplace=True)

    df['主观'] = df['主观'].round(0).astype(np.int64)
    df['客观'] = df['客观'].astype(np.int64)
    df['总分'] = df['总分'].astype(np.int64)
    # 选择数据行
    # df = df[df['等级'].isin(['优秀', '及格', '不及格'])]
    temp1 = df[:11]
    temp2 = df[-11:]
    df = pd.concat([temp1, temp2])
    # df = df.sort_values(by='总分', ascending=False)
    df.reset_index(drop=True, inplace=True)

    df = df[['序号', '节目名称', '频道', '播出时间', '录制地点',
             '制作方式', '制片人', '主观', '客观', '总分', '等级', '评语']]
    document = Document(file)
    # 将表格插入指定位置
    for p in document.paragraphs:
        if re.match("^Heading \d+$", p.style.name):
            if p.text == '专家意见及建议':
                print(p.text)
                # 插入表格
                table = document.add_table(rows=df.shape[0]*4+1, cols=df.shape[1]-1, style='Table Grid')
                # 移动表格到指定位置
                move_table_after(table, p)
    # 设置表格居中
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # 各列宽度
    table_width = {'序号': 8.5, '节目名称': 30, '频道': 20, '播出时间': 23.8, '录制地点': 21.6,
                   '制作方式': 18, '制片人': 15.8, '主观': 8.9, '客观': 9.7, '总分': 9.7, '等级': 12.3}
    # 取得各列名称
    columns_name = df.columns.to_list()
    columns_name.pop(columns_name.index('评语'))
    for i in columns_name:
        # 设置表格列宽
        table.columns[columns_name.index(i)].width = Mm(table_width[i])
        # 取得表格单元格
        cell1 = table.cell(0, columns_name.index(i))
        # 写入列名称
        cell1.text = i
        # 设置居中
        cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 设置垂直居中
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 设置表头底色
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value='#8DB4E2'))
        cell1._tc.get_or_add_tcPr().append(shading_elm_1)
        # 设置字体大小
        for run in cell1.paragraphs[0].runs:
            font = run.font
            font.size = Pt(10)
    # 写入数据
    table.rows[0].height = Mm(7.5)  # 表头行高
    for index, row in df.iterrows():
        for j in range(1):
            table.rows[index*4+j+1].height = Mm(7.5)  # 数据行高
            print(index*4+j+1)
        for i in range(len(row)-1):
            # 写入节目数据
            cell1 = table.cell(index * 4 + 1, i)
            cell1.text = str(row[columns_name[i]])
            cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # 设置字体大小
            for run in cell1.paragraphs[0].runs:
                font = run.font
                font.size = Pt(10)
        # 合并序号单元格
        cell1 = table.cell(index * 4 + 1, 0).merge(table.cell(index * 4 + 4, 0))
        # cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 合并评语单元格
        cell1 = table.cell(index * 4 + 2, 1).merge(table.cell(index * 4 + 4, len(row) - 2))
        # cell1.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell1.text = df.loc[index, '评语']
        cell1.paragraphs[0].paragraph_format.first_line_indent = Mm(7.4)
        # 设置字体大小
        for run in cell1.paragraphs[0].runs:
            font = run.font
            font.size = Pt(10)

    '''
    #设置表格宋体大小
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(10)
    '''
    document.save(file)

def write_to_Excel(file='database.xlsx', sheet_name='sheet1', start_row=0, start_col=0, df=pd.DataFrame[]):
    book = load_workbook(file)
    with pd.ExcelWriter(file,engine='openpyxl',datetime_format='%Y/%M/%D') as writer:
        writer.book = book
        writer.sheets = dict((ws.title, ws) for ws in book.worksheets)

        df.to_excel(writer, sheet_name, index=False, startrow=start_row, startcol=start_col)

if __name__ == '__main__':
    # initBaogao()
    # fenxi_youxiu()
    # fenxi_dabiao()
    # canping_program()
    # zonghe_fen()
    # rank_pindao()
    # fenxi_pindao()
    # rank_didian()
    # fenxi_didian()
    # rank_fangshi()
    # fenxi_fangshi()
    # Experts_zongping()
    
